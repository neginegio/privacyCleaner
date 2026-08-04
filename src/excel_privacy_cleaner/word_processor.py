from __future__ import annotations

import csv
import hashlib
import hmac
import json
import os
import re
import secrets
import shutil
import tempfile
import zipfile
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Iterator

from docx import Document

from .excel_processor import AliasBook, ProcessingOptions, replacement_for
from .ginza_japanese import GinzaEntityDetector, WORD_NLP_CONFIDENCE, WORD_NLP_DETECTION_RULE
from .presidio_japanese import JapanesePresidioDetector, entity_label


WORD_SUPPORTED_EXTENSION = ".docx"
WORD_MACRO_EXTENSION = ".docm"

W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
R_NS = "http://schemas.openxmlformats.org/officeDocument/2006/relationships"
REL_NS = "http://schemas.openxmlformats.org/package/2006/relationships"
WORD_CANDIDATE_CATEGORIES = {"会社名", "氏名", "住所", "電話番号", "メールアドレス", "銀行名"}
COMPANY_DESIGNATORS = ("株式会社", "有限会社", "合同会社", "医療法人", "学校法人", "社会福祉法人")
COMPANY_SUFFIX_DESIGNATORS = ("株式会社", "有限会社", "合同会社")
WORD_NAME_CHARS = r"一-龯々〆ヵヶぁ-んァ-ヶーA-Za-z0-9"
WORD_REVIEW_CONFIDENCE_THRESHOLD = 0.75
WORD_AUDIT_SCHEMA = "word_audit_v1"


class UnsupportedWordDocumentError(RuntimeError):
    pass


@dataclass(frozen=True)
class WordTextRun:
    story_type: str
    section_index: int | None
    container_type: str
    paragraph_index: int
    run_index: int
    char_start: int
    char_end: int
    text: str
    paragraph_text: str
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    cell_paragraph_index: int | None = None
    header_footer_type: str | None = None
    part_name: str = ""
    element_path: str = ""
    hyperlink_url: str | None = None
    hyperlink_rel_id: str | None = None
    is_linked_to_previous: bool = False

    @property
    def location_id(self) -> str:
        values = [
            self.story_type,
            str(self.section_index),
            self.header_footer_type or "",
            self.container_type,
            str(self.table_index),
            str(self.row_index),
            str(self.cell_index),
            str(self.paragraph_index),
            str(self.cell_paragraph_index),
            str(self.run_index),
            str(self.char_start),
            str(self.char_end),
            self.part_name,
        ]
        return ":".join(values)


@dataclass(frozen=True)
class WordParagraphText:
    story_type: str
    section_index: int | None
    container_type: str
    paragraph_index: int
    text: str
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    cell_paragraph_index: int | None = None
    header_footer_type: str | None = None
    part_name: str = ""
    element_path: str = ""
    is_linked_to_previous: bool = False


@dataclass(frozen=True)
class WordHyperlink:
    story_type: str
    section_index: int | None
    container_type: str
    paragraph_index: int
    text: str
    url: str
    rel_id: str
    table_index: int | None = None
    row_index: int | None = None
    cell_index: int | None = None
    cell_paragraph_index: int | None = None
    header_footer_type: str | None = None
    part_name: str = ""
    element_path: str = ""


@dataclass(frozen=True)
class WordDocumentProperties:
    author: str
    last_modified_by: str
    title: str
    subject: str
    keywords: str
    category: str
    comments: str
    created: datetime | None
    modified: datetime | None
    content_status: str
    identifier: str
    language: str
    revision: int | None
    version: str


@dataclass(frozen=True)
class UnsupportedWordFeature:
    feature_type: str
    part_name: str
    count: int
    detail: str = ""


@dataclass(frozen=True)
class WordStructureInventory:
    source_path: Path
    source_sha256: str
    paragraphs: tuple[WordParagraphText, ...]
    runs: tuple[WordTextRun, ...]
    hyperlinks: tuple[WordHyperlink, ...]
    document_properties: WordDocumentProperties
    unsupported_features: tuple[UnsupportedWordFeature, ...]
    macro_parts: tuple[str, ...]


@dataclass(frozen=True)
class WordCandidate:
    candidate_id: str
    category: str
    story_type: str
    section_index: int | None
    container_type: str
    paragraph_index: int
    table_index: int | None
    row_index: int | None
    cell_index: int | None
    location_id: str
    char_start: int
    char_end: int
    text: str
    detection_rule: str
    confidence: float
    source: str
    affected_run_indices: tuple[int, ...] = ()
    context_before: str = ""
    context_after: str = ""
    cell_paragraph_index: int | None = None
    header_footer_type: str | None = None
    original_category: str = ""
    hyperlink_url: str | None = None
    property_name: str | None = None


@dataclass
class WordReplacementDecision:
    candidate: WordCandidate
    enabled: bool = True
    replacement: str = ""
    # True only when a reviewer looked at a review-required candidate and
    # explicitly decided to keep the original text (not convert it), as
    # distinct from simply never having been reviewed yet. Only this state
    # (not a bare `enabled=False`) satisfies the review-required guard in
    # convert() without enabling the candidate.
    excluded: bool = False


@dataclass(frozen=True)
class WordConversionResult:
    output_path: Path
    warnings: tuple[str, ...]
    converted_run_count: int
    converted_property_count: int
    skipped_hyperlink_target_count: int
    review_required_count: int
    csv_path: Path
    report_path: Path


WORD_CATEGORY_ALIAS_KIND = {
    "会社名": "company",
    "氏名": "name",
    "住所": "address",
    "電話番号": "phone",
    "メールアドレス": "email",
}


def default_replacement_for_candidate(candidate: WordCandidate, alias_book: AliasBook, options: ProcessingOptions) -> str:
    if candidate.category == "銀行名":
        return alias_book.get("financial_institution", candidate.text)
    kind = WORD_CATEGORY_ALIAS_KIND.get(candidate.category)
    if kind:
        return replacement_for(kind, candidate.text, alias_book, options)
    return alias_book.get("text", candidate.text)


def candidate_requires_review(candidate: WordCandidate) -> bool:
    if candidate.source == "hyperlink_target":
        return False
    return candidate.confidence < WORD_REVIEW_CONFIDENCE_THRESHOLD


def build_default_decisions(
    candidates: tuple[WordCandidate, ...],
    alias_book: AliasBook,
    options: ProcessingOptions,
) -> list[WordReplacementDecision]:
    decisions: list[WordReplacementDecision] = []
    for candidate in candidates:
        enabled = candidate.source != "hyperlink_target" and not candidate_requires_review(candidate)
        replacement = default_replacement_for_candidate(candidate, alias_book, options)
        decisions.append(WordReplacementDecision(candidate=candidate, enabled=enabled, replacement=replacement))
    return decisions


def extract_word_structure(path: Path) -> WordStructureInventory:
    validate_docx_input(path)
    source_sha256 = _file_sha256(path)
    unsupported_features, macro_parts = inspect_docx_package(path)
    if macro_parts:
        raise UnsupportedWordDocumentError("マクロを含むWord文書はv1対象外です: " + ", ".join(macro_parts))

    document = Document(path)
    paragraphs: list[WordParagraphText] = []
    runs: list[WordTextRun] = []
    hyperlinks: list[WordHyperlink] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        _append_paragraph(
            paragraphs,
            runs,
            hyperlinks,
            paragraph,
            story_type="body",
            section_index=None,
            container_type="paragraph",
            paragraph_index=paragraph_index,
            element_path=f"body/p[{paragraph_index}]",
        )

    for table_index, table in enumerate(document.tables):
        _append_table(
            paragraphs,
            runs,
            hyperlinks,
            table,
            story_type="body",
            section_index=None,
            header_footer_type=None,
            table_index=table_index,
            part_name=_part_name(table),
            element_path_prefix=f"body/tbl[{table_index}]",
        )

    for section_index, section in enumerate(document.sections):
        for header_footer_type, story_type, story in _section_stories(section):
            linked = bool(story.is_linked_to_previous)
            part_name = _part_name(story)
            story_prefix = f"{story_type}[{section_index}]/{header_footer_type}"
            for paragraph_index, paragraph in enumerate(story.paragraphs):
                _append_paragraph(
                    paragraphs,
                    runs,
                    hyperlinks,
                    paragraph,
                    story_type=story_type,
                    section_index=section_index,
                    container_type="paragraph",
                    paragraph_index=paragraph_index,
                    header_footer_type=header_footer_type,
                    part_name=part_name,
                    element_path=f"{story_prefix}/p[{paragraph_index}]",
                    is_linked_to_previous=linked,
                )
            for table_index, table in enumerate(story.tables):
                _append_table(
                    paragraphs,
                    runs,
                    hyperlinks,
                    table,
                    story_type=story_type,
                    section_index=section_index,
                    header_footer_type=header_footer_type,
                    table_index=table_index,
                    part_name=part_name or _part_name(table),
                    element_path_prefix=f"{story_prefix}/tbl[{table_index}]",
                    is_linked_to_previous=linked,
                )

    return WordStructureInventory(
        source_path=path,
        source_sha256=source_sha256,
        paragraphs=tuple(paragraphs),
        runs=tuple(runs),
        hyperlinks=tuple(hyperlinks),
        document_properties=_document_properties(document),
        unsupported_features=unsupported_features,
        macro_parts=macro_parts,
    )


def _section_stories(section: Any) -> tuple[tuple[str, str, Any], ...]:
    return (
        ("default", "header", section.header),
        ("first_page", "header", section.first_page_header),
        ("even_page", "header", section.even_page_header),
        ("default", "footer", section.footer),
        ("first_page", "footer", section.first_page_footer),
        ("even_page", "footer", section.even_page_footer),
    )


def _iter_table_paragraph_objects(table: Any) -> Iterator[Any]:
    for row in table.rows:
        for cell in row.cells:
            yield from cell.paragraphs


def _iter_paragraph_objects(document: Any) -> Iterator[Any]:
    yield from document.paragraphs
    for table in document.tables:
        yield from _iter_table_paragraph_objects(table)
    for section in document.sections:
        for _header_footer_type, _story_type, story in _section_stories(section):
            yield from story.paragraphs
            for table in story.tables:
                yield from _iter_table_paragraph_objects(table)


def detect_word_candidates(path: Path) -> tuple[WordStructureInventory, tuple[WordCandidate, ...]]:
    inventory = extract_word_structure(path)
    return inventory, candidates_for_inventory(inventory)


def candidates_for_inventory(inventory: WordStructureInventory) -> tuple[WordCandidate, ...]:
    detector = JapanesePresidioDetector()
    ginza_detector = GinzaEntityDetector()
    candidates: list[WordCandidate] = []
    seen: set[tuple[str, str, int, int, str, str]] = set()
    runs_by_location = _runs_by_paragraph_location(inventory.runs)

    for paragraph in inventory.paragraphs:
        if not paragraph.text.strip():
            continue
        runs = runs_by_location.get(_paragraph_location_id(paragraph), ())
        _append_text_candidates(candidates, seen, inventory, paragraph, paragraph.text, runs, detector, ginza_detector, source="paragraph")

    for hyperlink in inventory.hyperlinks:
        if hyperlink.url:
            _append_hyperlink_target_candidates(candidates, seen, inventory, hyperlink)

    for property_name, value in _document_property_values(inventory.document_properties):
        if value:
            _append_property_candidates(candidates, seen, inventory, property_name, value, detector, ginza_detector)

    return tuple(candidates)


class WordPrivacyProcessor:
    def __init__(self) -> None:
        self.detector = JapanesePresidioDetector()
        self.alias_book = AliasBook()
        self.options = ProcessingOptions()
        self.inventory: WordStructureInventory | None = None
        self.temp_dir: Path | None = None
        self.temp_docx: Path | None = None

    def cleanup(self) -> None:
        if self.temp_dir and self.temp_dir.exists():
            shutil.rmtree(self.temp_dir, ignore_errors=True)
        self.temp_dir = None
        self.temp_docx = None
        self.inventory = None

    def scan(self, docx_path: Path, options: ProcessingOptions | None = None) -> list[WordReplacementDecision]:
        self.cleanup()
        self.options = options or ProcessingOptions()
        self.alias_book = AliasBook()
        self.temp_dir = Path(tempfile.mkdtemp(prefix="WordPrivacyCleaner_"))
        self.temp_docx = self.temp_dir / docx_path.name
        shutil.copy2(docx_path, self.temp_docx)
        self.inventory = extract_word_structure(self.temp_docx)
        candidates = candidates_for_inventory(self.inventory)
        return build_default_decisions(candidates, self.alias_book, self.options)

    def convert(
        self,
        source_path: Path,
        decisions: list[WordReplacementDecision],
        output_dir: Path | None = None,
        write_artifacts: bool = True,
    ) -> WordConversionResult:
        if not self.temp_docx or not self.temp_docx.exists() or self.inventory is None:
            raise RuntimeError("先に検査を実行してください。")

        warnings: list[str] = []
        enabled_decisions = [decision for decision in decisions if decision.enabled]

        blocked_hyperlink_targets = [
            decision for decision in enabled_decisions if decision.candidate.source == "hyperlink_target"
        ]
        if blocked_hyperlink_targets:
            preview = "、".join(decision.candidate.text for decision in blocked_hyperlink_targets[:8])
            raise RuntimeError(
                f"ハイパーリンクURLの置換は未対応です。対象候補 {len(blocked_hyperlink_targets)} 件を無効のままにしてください。"
                f"対象例: {preview}"
            )
        skipped_hyperlink_target_count = sum(
            1 for decision in decisions if decision.candidate.source == "hyperlink_target" and not decision.enabled
        )
        if skipped_hyperlink_target_count:
            warnings.append(
                f"ハイパーリンクURL候補 {skipped_hyperlink_target_count} 件は今回の匿名化対象外です(URL自体は元のまま残ります)。"
            )

        review_required = [
            decision
            for decision in decisions
            if not decision.enabled
            and not decision.excluded
            and decision.candidate.source != "hyperlink_target"
            and candidate_requires_review(decision.candidate)
        ]
        if review_required:
            preview = "、".join(f"{decision.candidate.category}:{decision.candidate.text}" for decision in review_required[:8])
            raise RuntimeError(
                f"未確認候補が {len(review_required)} 件あります。信頼度の低い検出のため内容を確認し、"
                f"問題なければ候補を有効にするか、誤りであれば手動修正のうえ再実行してください。対象例: {preview}"
            )

        if self.inventory.unsupported_features:
            feature_counts: dict[str, int] = {}
            for feature in self.inventory.unsupported_features:
                feature_counts[feature.feature_type] = feature_counts.get(feature.feature_type, 0) + feature.count
            detail = "、".join(f"{feature_type}: {count}件" for feature_type, count in sorted(feature_counts.items()))
            if not self.options.is_analysis:
                raise RuntimeError(f"未対応領域が検出されたため外部共有用の出力を停止しました。未対応領域: {detail}")
            warnings.append(
                f"未対応領域が検出されました(未対応領域: {detail})。未対応領域内の匿名化は保証しない。"
                "分析継続用の範囲で利用し、外部共有する場合は内容を必ず確認してください。"
            )

        paragraph_by_location = {
            _paragraph_location_id(paragraph): paragraph for paragraph in self.inventory.paragraphs
        }
        paragraph_decisions = [decision for decision in enabled_decisions if decision.candidate.source == "paragraph"]
        for decision in paragraph_decisions:
            candidate = decision.candidate
            paragraph = paragraph_by_location.get(candidate.location_id)
            if paragraph is None or paragraph.text[candidate.char_start : candidate.char_end] != candidate.text:
                raise RuntimeError(
                    f"構造が変化しました。再スキャンしてください。対象候補: {candidate.category} 「{candidate.text}」"
                )

        decisions_by_location: dict[str, list[WordReplacementDecision]] = {}
        for decision in paragraph_decisions:
            decisions_by_location.setdefault(decision.candidate.location_id, []).append(decision)
        for location_id, location_decisions in decisions_by_location.items():
            ordered = sorted(location_decisions, key=lambda item: item.candidate.char_start)
            for previous, current in zip(ordered, ordered[1:]):
                if previous.candidate.char_end > current.candidate.char_start:
                    raise RuntimeError(
                        "検出候補の範囲が重複しています: "
                        f"{previous.candidate.category} 「{previous.candidate.text}」 / "
                        f"{current.candidate.category} 「{current.candidate.text}」"
                    )

        property_decisions = [decision for decision in enabled_decisions if decision.candidate.source == "document_property"]
        decisions_by_property: dict[str, list[WordReplacementDecision]] = {}
        for decision in property_decisions:
            property_name = decision.candidate.property_name or ""
            decisions_by_property.setdefault(property_name, []).append(decision)
        for property_name, property_decision_list in decisions_by_property.items():
            ordered = sorted(property_decision_list, key=lambda item: item.candidate.char_start)
            for previous, current in zip(ordered, ordered[1:]):
                if previous.candidate.char_end > current.candidate.char_start:
                    raise RuntimeError(
                        f"文書プロパティ {property_name} の検出候補の範囲が重複しています。"
                    )

        document = Document(self.temp_docx)
        live_paragraph_objects = list(_iter_paragraph_objects(document))
        if len(live_paragraph_objects) != len(self.inventory.paragraphs):
            raise RuntimeError("構造が変化しました。再スキャンしてください。")
        live_paragraphs_by_location = {
            _paragraph_location_id(paragraph_info): paragraph_object
            for paragraph_info, paragraph_object in zip(self.inventory.paragraphs, live_paragraph_objects)
        }

        converted_run_count = 0
        for location_id, location_decisions in decisions_by_location.items():
            paragraph_object = live_paragraphs_by_location.get(location_id)
            if paragraph_object is None:
                continue
            run_items = _paragraph_runs_with_offsets(paragraph_object)
            converted_run_count += _apply_paragraph_decisions(run_items, location_decisions)

        converted_property_count = 0
        for property_name, property_decision_list in decisions_by_property.items():
            current_value = getattr(document.core_properties, property_name, "") or ""
            new_value = current_value
            for start, end, replacement in sorted(
                (
                    (decision.candidate.char_start, decision.candidate.char_end, decision.replacement)
                    for decision in property_decision_list
                ),
                key=lambda item: item[0],
                reverse=True,
            ):
                new_value = new_value[:start] + replacement + new_value[end:]
            setattr(document.core_properties, property_name, new_value)
            converted_property_count += 1

        output_path = _make_word_output_path(source_path, output_dir=output_dir, options=self.options)
        document.save(output_path)

        residual = _find_residual_text(output_path, paragraph_decisions + property_decisions)
        if residual:
            detail = _format_residual_detail(residual)
            if not self.options.is_analysis:
                output_path.unlink(missing_ok=True)
                raise RuntimeError(
                    f"内部XML残存検査で匿名化対象の原文が検出されたため外部共有用の出力を停止しました。残存箇所: {detail}"
                )
            warnings.append(
                f"内部XML残存検査で匿名化対象の原文が検出されました(残存箇所: {detail})。"
                "本文以外の内部XMLに原文が残っている可能性があるため、外部共有する場合は内容を必ず確認してください。"
            )

        review_required_count = sum(
            1
            for decision in decisions
            if decision.candidate.source != "hyperlink_target" and candidate_requires_review(decision.candidate)
        )
        csv_path = output_path.with_name(f"{output_path.stem}_検出変換結果.csv")
        report_path = output_path.with_name(f"{output_path.stem}_処理報告書.txt")
        result = WordConversionResult(
            output_path=output_path,
            warnings=tuple(warnings),
            converted_run_count=converted_run_count,
            converted_property_count=converted_property_count,
            skipped_hyperlink_target_count=skipped_hyperlink_target_count,
            review_required_count=review_required_count,
            csv_path=csv_path,
            report_path=report_path,
        )
        if write_artifacts:
            audit_path = output_path.with_name(f"{output_path.stem}_監査報告.json")
            write_word_findings_csv(csv_path, decisions)
            write_word_processing_report(
                report_path,
                source_path=source_path,
                result=result,
                decisions=decisions,
                inventory=self.inventory,
                options=self.options,
            )
            write_word_audit_json(
                audit_path,
                source_path=source_path,
                result=result,
                decisions=decisions,
                inventory=self.inventory,
                options=self.options,
            )
        self.cleanup()
        return result


def _apply_paragraph_decisions(run_items: list[dict[str, Any]], decisions: list[WordReplacementDecision]) -> int:
    edits_by_run: dict[int, list[tuple[int, int, str]]] = {}
    for decision in decisions:
        candidate = decision.candidate
        for position, run_index in enumerate(candidate.affected_run_indices):
            if run_index >= len(run_items):
                continue
            run_item = run_items[run_index]
            local_start = max(candidate.char_start, run_item["char_start"]) - run_item["char_start"]
            local_end = min(candidate.char_end, run_item["char_end"]) - run_item["char_start"]
            slice_text = decision.replacement if position == 0 else ""
            edits_by_run.setdefault(run_index, []).append((local_start, local_end, slice_text))

    changed = 0
    for run_index, edits in edits_by_run.items():
        text = run_items[run_index]["text"]
        for start, end, slice_text in sorted(edits, key=lambda item: item[0], reverse=True):
            text = text[:start] + slice_text + text[end:]
        run_items[run_index]["element"].text = text
        changed += 1
    return changed


def _find_residual_text(
    output_path: Path,
    checked_decisions: list[WordReplacementDecision],
) -> dict[str, set[str]]:
    """Scan every .xml/.rels part of the saved output package for leftover
    original text from checked_decisions. Callers must pre-filter to
    paragraph/document_property decisions only -- hyperlink_target candidates
    must never be passed in, since their original text (inside a URL) is
    deliberately left unconverted by design and is not a residual leak.
    """
    residual: dict[str, set[str]] = {}
    texts_by_category: dict[str, set[str]] = {}
    for decision in checked_decisions:
        text = decision.candidate.text
        if text:
            texts_by_category.setdefault(decision.candidate.category, set()).add(text)
    if not texts_by_category:
        return residual

    with zipfile.ZipFile(output_path) as archive:
        for name in archive.namelist():
            if not name.endswith((".xml", ".rels")):
                continue
            content = archive.read(name).decode("utf-8", errors="ignore")
            found = {
                category
                for category, texts in texts_by_category.items()
                if any(text in content for text in texts)
            }
            if found:
                residual[name] = found
    return residual


def _format_residual_detail(residual: dict[str, set[str]]) -> str:
    return "、".join(
        f"{part_name}: {', '.join(sorted(categories))}"
        for part_name, categories in sorted(residual.items())
    )


def word_candidate_location_label(candidate: WordCandidate) -> str:
    if candidate.story_type == "document_property":
        return f"文書プロパティ({candidate.property_name or ''})"

    if candidate.story_type == "body":
        story_label = "本文"
    elif candidate.story_type == "header":
        story_label = f"ヘッダー({candidate.header_footer_type or 'default'})"
    elif candidate.story_type == "footer":
        story_label = f"フッター({candidate.header_footer_type or 'default'})"
    else:
        story_label = candidate.story_type

    if candidate.section_index is not None and candidate.section_index > 0:
        story_label += f" セクション{candidate.section_index + 1}"

    if candidate.container_type == "table_cell":
        location = (
            f"{story_label} 表{(candidate.table_index or 0) + 1} "
            f"行{(candidate.row_index or 0) + 1}列{(candidate.cell_index or 0) + 1}"
        )
        if candidate.cell_paragraph_index:
            location += f" 段落{candidate.cell_paragraph_index + 1}"
        return location

    return f"{story_label} 段落{candidate.paragraph_index + 1}"


def word_finding_status(decision: WordReplacementDecision) -> str:
    candidate = decision.candidate
    if candidate.source == "hyperlink_target":
        return "ハイパーリンク対象外"
    requires_review = candidate_requires_review(candidate)
    if decision.enabled:
        return "要確認・承認済み" if requires_review else "自動変換"
    if decision.excluded:
        return "確認済み(除外)"
    return "要確認(未処理)" if requires_review else "維持(手動)"


def word_finding_reason(decision: WordReplacementDecision) -> str:
    candidate = decision.candidate
    base = f"{candidate.detection_rule}(信頼度{candidate.confidence:.2f})"
    suffix = {
        "ハイパーリンク対象外": " / ハイパーリンクURLの置換は未対応のため対象外",
        "要確認・承認済み": " / 要確認候補を利用者が確認のうえ承認",
        "要確認(未処理)": " / 要確認候補が未承認のため対象外",
        "確認済み(除外)": " / 要確認候補を利用者が確認のうえ除外(原文を維持)",
        "維持(手動)": " / 利用者が対象を解除し原文を維持",
        "自動変換": "",
    }[word_finding_status(decision)]
    return base + suffix


def _make_word_output_path(source_path: Path, output_dir: Path | None = None, options: ProcessingOptions | None = None) -> Path:
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    folder = output_dir if output_dir is not None else source_path.parent
    folder.mkdir(parents=True, exist_ok=True)
    mode_label = (options or ProcessingOptions()).mode_label
    candidate = folder / f"{source_path.stem}_{mode_label}_匿名化_{timestamp}{WORD_SUPPORTED_EXTENSION}"
    if not candidate.exists():
        return candidate
    return folder / f"{source_path.stem}_{mode_label}_匿名化_{timestamp}_{datetime.now().microsecond:06d}{WORD_SUPPORTED_EXTENSION}"


def write_word_findings_csv(path: Path, decisions: list[WordReplacementDecision]) -> None:
    headers = ["変換", "処理状態", "位置", "種類", "検査", "検出値", "変換後", "理由"]
    with path.open("w", encoding="utf-8-sig", newline="") as output:
        writer = csv.writer(output)
        writer.writerow(headers)
        for decision in decisions:
            candidate = decision.candidate
            writer.writerow(
                [
                    "対象" if decision.enabled else "維持",
                    word_finding_status(decision),
                    word_candidate_location_label(candidate),
                    candidate.category,
                    candidate.detection_rule,
                    candidate.text,
                    decision.replacement,
                    word_finding_reason(decision),
                ]
            )


def write_word_processing_report(
    path: Path,
    *,
    source_path: Path,
    result: WordConversionResult,
    decisions: list[WordReplacementDecision],
    inventory: WordStructureInventory,
    options: ProcessingOptions,
) -> None:
    lines = [
        "Word匿名化 処理報告書",
        "",
        f"処理モード: {options.mode_label}",
        f"入力ファイル名: {source_path.name}",
        f"出力ファイル名: {result.output_path.name}",
        f"処理日時: {datetime.now():%Y-%m-%d %H:%M:%S}",
        f"検出候補数: {len(decisions)}",
        f"変換run数: {result.converted_run_count}",
        f"変換文書プロパティ数: {result.converted_property_count}",
        f"要確認件数: {result.review_required_count}",
        f"ハイパーリンクURL対象外件数: {result.skipped_hyperlink_target_count}",
        "",
        "未対応領域:",
    ]
    if inventory.unsupported_features:
        lines.extend(
            f"- {feature.feature_type} ({feature.part_name}): {feature.count}件"
            for feature in inventory.unsupported_features
        )
    else:
        lines.append("- なし")
    lines.extend(["", "警告内容:"])
    lines.extend([f"- {warning}" for warning in result.warnings] or ["- なし"])
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def write_word_audit_json(
    path: Path,
    *,
    source_path: Path,
    result: WordConversionResult,
    decisions: list[WordReplacementDecision],
    inventory: WordStructureInventory,
    options: ProcessingOptions,
) -> None:
    payload = {
        "schema": WORD_AUDIT_SCHEMA,
        "summary": {
            "input_file": source_path.name,
            "output_file": result.output_path.name,
            "mode": options.mode_label,
            "processed_at": datetime.now().isoformat(timespec="seconds"),
            "source_sha256": inventory.source_sha256,
            "candidate_count": len(decisions),
            "converted_run_count": result.converted_run_count,
            "converted_property_count": result.converted_property_count,
            "review_required_count": result.review_required_count,
            "skipped_hyperlink_target_count": result.skipped_hyperlink_target_count,
            "unsupported_features": [
                {"feature_type": feature.feature_type, "part_name": feature.part_name, "count": feature.count}
                for feature in inventory.unsupported_features
            ],
            "warnings": list(result.warnings),
        },
        "findings": [
            {
                "category": decision.candidate.category,
                "location": word_candidate_location_label(decision.candidate),
                "detection_rule": decision.candidate.detection_rule,
                "confidence": decision.candidate.confidence,
                "source": decision.candidate.source,
                "status": word_finding_status(decision),
                "enabled": decision.enabled,
                "replacement": decision.replacement,
                "reason": word_finding_reason(decision),
                "original_hmac_sha256": _word_hmac_digest(decision.candidate.text) if decision.candidate.text else "",
            }
            for decision in decisions
        ],
    }
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def validate_docx_input(path: Path) -> None:
    suffix = path.suffix.lower()
    if suffix == WORD_MACRO_EXTENSION:
        raise UnsupportedWordDocumentError(".docmはWord匿名化v1の対象外です。")
    if suffix != WORD_SUPPORTED_EXTENSION:
        raise UnsupportedWordDocumentError(f"Word匿名化v1は.docxのみ対応です: {path.name}")
    if not path.exists():
        raise FileNotFoundError(path)


def inspect_docx_package(path: Path) -> tuple[tuple[UnsupportedWordFeature, ...], tuple[str, ...]]:
    features: dict[tuple[str, str], int] = {}
    details: dict[tuple[str, str], str] = {}
    macro_parts: list[str] = []

    try:
        with zipfile.ZipFile(path) as archive:
            names = archive.namelist()
            for name in names:
                lowered = name.lower()
                if lowered.endswith("vbaproject.bin") or "vbaproject" in lowered:
                    macro_parts.append(name)
                if name.startswith("word/embeddings/"):
                    _count_feature(features, details, "embedded_object", name)
                if name.startswith("customXml/"):
                    _count_feature(features, details, "custom_xml", name)
                if name.startswith("word/comments") and name.endswith(".xml"):
                    _count_feature(features, details, "comments", name)
                if name == "word/footnotes.xml":
                    _count_feature(features, details, "footnotes", name)
                if name == "word/endnotes.xml":
                    _count_feature(features, details, "endnotes", name)
                if not name.endswith((".xml", ".rels")):
                    continue
                data = archive.read(name)
                text = data.decode("utf-8", errors="ignore")
                _inspect_xml_text(name, text, features, details)
    except zipfile.BadZipFile as exc:
        raise UnsupportedWordDocumentError(f"DOCX ZIP構造を読み取れませんでした: {exc}") from exc

    unsupported = tuple(
        UnsupportedWordFeature(feature_type=feature_type, part_name=part_name, count=count, detail=details.get((feature_type, part_name), ""))
        for (feature_type, part_name), count in sorted(features.items())
    )
    return unsupported, tuple(sorted(macro_parts))


def _append_table(
    paragraphs: list[WordParagraphText],
    runs: list[WordTextRun],
    hyperlinks: list[WordHyperlink],
    table: Any,
    *,
    story_type: str,
    section_index: int | None,
    header_footer_type: str | None,
    table_index: int,
    part_name: str,
    element_path_prefix: str,
    is_linked_to_previous: bool = False,
) -> None:
    for row_index, row in enumerate(table.rows):
        for cell_index, cell in enumerate(row.cells):
            for cell_paragraph_index, paragraph in enumerate(cell.paragraphs):
                _append_paragraph(
                    paragraphs,
                    runs,
                    hyperlinks,
                    paragraph,
                    story_type=story_type,
                    section_index=section_index,
                    container_type="table_cell",
                    paragraph_index=cell_paragraph_index,
                    table_index=table_index,
                    row_index=row_index,
                    cell_index=cell_index,
                    cell_paragraph_index=cell_paragraph_index,
                    header_footer_type=header_footer_type,
                    part_name=part_name,
                    element_path=f"{element_path_prefix}/tr[{row_index}]/tc[{cell_index}]/p[{cell_paragraph_index}]",
                    is_linked_to_previous=is_linked_to_previous,
                )


def _append_paragraph(
    paragraphs: list[WordParagraphText],
    runs: list[WordTextRun],
    hyperlinks: list[WordHyperlink],
    paragraph: Any,
    *,
    story_type: str,
    section_index: int | None,
    container_type: str,
    paragraph_index: int,
    table_index: int | None = None,
    row_index: int | None = None,
    cell_index: int | None = None,
    cell_paragraph_index: int | None = None,
    header_footer_type: str | None = None,
    part_name: str = "",
    element_path: str,
    is_linked_to_previous: bool = False,
) -> None:
    part_name = part_name or _part_name(paragraph)
    extracted_runs = _paragraph_runs_with_offsets(paragraph)
    paragraph_text = "".join(item["text"] for item in extracted_runs)
    paragraphs.append(
        WordParagraphText(
            story_type=story_type,
            section_index=section_index,
            container_type=container_type,
            paragraph_index=paragraph_index,
            text=paragraph_text,
            table_index=table_index,
            row_index=row_index,
            cell_index=cell_index,
            cell_paragraph_index=cell_paragraph_index,
            header_footer_type=header_footer_type,
            part_name=part_name,
            element_path=element_path,
            is_linked_to_previous=is_linked_to_previous,
        )
    )
    for run_index, item in enumerate(extracted_runs):
        text = str(item["text"])
        run_path = f"{element_path}/r[{run_index}]"
        run = WordTextRun(
            story_type=story_type,
            section_index=section_index,
            container_type=container_type,
            paragraph_index=paragraph_index,
            table_index=table_index,
            row_index=row_index,
            cell_index=cell_index,
            cell_paragraph_index=cell_paragraph_index,
            run_index=run_index,
            char_start=int(item["char_start"]),
            char_end=int(item["char_end"]),
            text=text,
            paragraph_text=paragraph_text,
            header_footer_type=header_footer_type,
            part_name=part_name,
            element_path=run_path,
            hyperlink_url=item.get("hyperlink_url"),
            hyperlink_rel_id=item.get("hyperlink_rel_id"),
            is_linked_to_previous=is_linked_to_previous,
        )
        runs.append(run)
    for item in extracted_runs:
        rel_id = item.get("hyperlink_rel_id")
        url = item.get("hyperlink_url")
        if rel_id and url:
            hyperlinks.append(
                WordHyperlink(
                    story_type=story_type,
                    section_index=section_index,
                    container_type=container_type,
                    paragraph_index=paragraph_index,
                    table_index=table_index,
                    row_index=row_index,
                    cell_index=cell_index,
                    cell_paragraph_index=cell_paragraph_index,
                    header_footer_type=header_footer_type,
                    text=str(item["text"]),
                    url=str(url),
                    rel_id=str(rel_id),
                    part_name=part_name,
                    element_path=element_path,
                )
            )


def _paragraph_runs_with_offsets(paragraph: Any) -> list[dict[str, Any]]:
    items: list[dict[str, Any]] = []
    cursor = 0
    for child in paragraph._p.iterchildren():
        tag = _local_name(child.tag)
        if tag == "r":
            text = _run_text(child)
            if text:
                items.append({"text": text, "char_start": cursor, "char_end": cursor + len(text), "element": child})
                cursor += len(text)
        elif tag == "hyperlink":
            rel_id = child.get(f"{{{R_NS}}}id")
            url = _hyperlink_target(paragraph, rel_id) if rel_id else None
            for run in child.iterchildren():
                if _local_name(run.tag) != "r":
                    continue
                text = _run_text(run)
                if text:
                    items.append(
                        {
                            "text": text,
                            "char_start": cursor,
                            "char_end": cursor + len(text),
                            "hyperlink_url": url,
                            "hyperlink_rel_id": rel_id,
                            "element": run,
                        }
                    )
                    cursor += len(text)
    return items


def _run_text(run_element: Any) -> str:
    parts: list[str] = []
    for child in run_element.iterchildren():
        tag = _local_name(child.tag)
        if tag == "t":
            parts.append(child.text or "")
        elif tag == "tab":
            parts.append("\t")
        elif tag in {"br", "cr"}:
            parts.append("\n")
    return "".join(parts)


def _hyperlink_target(paragraph: Any, rel_id: str | None) -> str | None:
    if not rel_id:
        return None
    part = getattr(paragraph, "part", None)
    if part is None:
        parent = getattr(paragraph, "_parent", None)
        part = getattr(parent, "part", None)
    if part is None:
        return None
    relationship = part.rels.get(rel_id)
    return getattr(relationship, "target_ref", None) if relationship is not None else None


def _document_properties(document: Any) -> WordDocumentProperties:
    props = document.core_properties
    return WordDocumentProperties(
        author=str(props.author or ""),
        last_modified_by=str(props.last_modified_by or ""),
        title=str(props.title or ""),
        subject=str(props.subject or ""),
        keywords=str(props.keywords or ""),
        category=str(props.category or ""),
        comments=str(props.comments or ""),
        created=props.created,
        modified=props.modified,
        content_status=str(props.content_status or ""),
        identifier=str(props.identifier or ""),
        language=str(props.language or ""),
        revision=props.revision,
        version=str(props.version or ""),
    )


def _append_text_candidates(
    candidates: list[WordCandidate],
    seen: set[tuple[str, str, int, int, str, str]],
    inventory: WordStructureInventory,
    paragraph: WordParagraphText,
    text: str,
    runs: tuple[WordTextRun, ...],
    detector: JapanesePresidioDetector,
    ginza_detector: GinzaEntityDetector | None,
    *,
    source: str,
) -> None:
    for start, end, category, rule, confidence, original_category in _word_detector_results(text, detector, ginza_detector):
        _append_candidate(
            candidates,
            seen,
            inventory,
            paragraph,
            text,
            start,
            end,
            category,
            rule,
            confidence,
            source,
            affected_runs=_affected_run_indices(runs, start, end),
            original_category=original_category,
        )


def _append_hyperlink_target_candidates(
    candidates: list[WordCandidate],
    seen: set[tuple[str, str, int, int, str, str]],
    inventory: WordStructureInventory,
    hyperlink: WordHyperlink,
) -> None:
    paragraph = WordParagraphText(
        story_type=hyperlink.story_type,
        section_index=hyperlink.section_index,
        container_type=hyperlink.container_type,
        paragraph_index=hyperlink.paragraph_index,
        table_index=hyperlink.table_index,
        row_index=hyperlink.row_index,
        cell_index=hyperlink.cell_index,
        cell_paragraph_index=hyperlink.cell_paragraph_index,
        header_footer_type=hyperlink.header_footer_type,
        text=hyperlink.url,
        part_name=hyperlink.part_name,
        element_path=hyperlink.element_path,
    )
    for start, end, category, rule, confidence in _url_candidate_results(hyperlink.url):
        _append_candidate(
            candidates,
            seen,
            inventory,
            paragraph,
            hyperlink.url,
            start,
            end,
            category,
            rule,
            confidence,
            "hyperlink_target",
            affected_runs=(),
            original_category=category,
            hyperlink_url=hyperlink.url,
        )


def _append_property_candidates(
    candidates: list[WordCandidate],
    seen: set[tuple[str, str, int, int, str, str]],
    inventory: WordStructureInventory,
    property_name: str,
    value: str,
    detector: JapanesePresidioDetector,
    ginza_detector: GinzaEntityDetector | None,
) -> None:
    paragraph = WordParagraphText(
        story_type="document_property",
        section_index=None,
        container_type="property",
        paragraph_index=0,
        text=value,
        element_path=f"docProps/core.xml/{property_name}",
        part_name="/docProps/core.xml",
    )
    for start, end, category, rule, confidence, original_category in _word_detector_results(value, detector, ginza_detector):
        _append_candidate(
            candidates,
            seen,
            inventory,
            paragraph,
            value,
            start,
            end,
            category,
            rule,
            confidence,
            "document_property",
            affected_runs=(),
            original_category=original_category,
            property_name=property_name,
        )


def _append_candidate(
    candidates: list[WordCandidate],
    seen: set[tuple[str, str, int, int, str, str]],
    inventory: WordStructureInventory,
    paragraph: WordParagraphText,
    source_text: str,
    start: int,
    end: int,
    category: str,
    detection_rule: str,
    confidence: float,
    source: str,
    *,
    affected_runs: tuple[int, ...],
    original_category: str,
    hyperlink_url: str | None = None,
    property_name: str | None = None,
) -> None:
    if start < 0 or end <= start or end > len(source_text):
        return
    value = source_text[start:end]
    if not value.strip():
        return
    location_id = _paragraph_location_id(paragraph)
    key = (location_id, category, start, end, value, source)
    if key in seen:
        return
    seen.add(key)
    candidates.append(
        WordCandidate(
            candidate_id=_candidate_id(inventory.source_sha256, location_id, category, start, end, value, source),
            category=category,
            story_type=paragraph.story_type,
            section_index=paragraph.section_index,
            container_type=paragraph.container_type,
            paragraph_index=paragraph.paragraph_index,
            table_index=paragraph.table_index,
            row_index=paragraph.row_index,
            cell_index=paragraph.cell_index,
            cell_paragraph_index=paragraph.cell_paragraph_index,
            header_footer_type=paragraph.header_footer_type,
            location_id=location_id,
            char_start=start,
            char_end=end,
            text=value,
            detection_rule=detection_rule,
            confidence=confidence,
            source=source,
            affected_run_indices=affected_runs,
            context_before=source_text[max(0, start - 12) : start],
            context_after=source_text[end : min(len(source_text), end + 12)],
            original_category=original_category,
            hyperlink_url=hyperlink_url,
            property_name=property_name,
        )
    )


def _word_detector_results(
    text: str,
    detector: JapanesePresidioDetector,
    ginza_detector: GinzaEntityDetector | None = None,
) -> list[tuple[int, int, str, str, float, str]]:
    results: list[tuple[int, int, str, str, float, str]] = []
    for detector_result in detector.analyze(text):
        label = _normalize_word_category(entity_label(detector_result.entity_type))
        if label not in WORD_CANDIDATE_CATEGORIES:
            continue
        start, end = _trim_candidate_span(text, detector_result.start, detector_result.end, label)
        if label == "住所":
            start, end = _extend_address_span(text, start, end)
            if _is_weak_address_candidate(text[start:end]):
                continue
        if _is_generic_word_false_positive(label, text[start:end]):
            continue
        results.append((start, end, label, "JapanesePresidioDetector", float(detector_result.score), detector_result.entity_type))
    results.extend(_regex_candidate_results(text))
    if ginza_detector is not None:
        results.extend(_ginza_candidate_results(text, ginza_detector))
    return _select_non_overlapping_word_results(text, results)


def _ginza_candidate_results(
    text: str,
    ginza_detector: GinzaEntityDetector,
) -> list[tuple[int, int, str, str, float, str]]:
    results: list[tuple[int, int, str, str, float, str]] = []
    for nlp_result in ginza_detector.analyze(text):
        start, end = _trim_candidate_span(text, nlp_result.start, nlp_result.end, nlp_result.category)
        value = text[start:end]
        if _is_generic_word_false_positive(nlp_result.category, value):
            continue
        results.append((start, end, nlp_result.category, WORD_NLP_DETECTION_RULE, WORD_NLP_CONFIDENCE, nlp_result.raw_label))
    return results


def _regex_candidate_results(text: str) -> list[tuple[int, int, str, str, float, str]]:
    results: list[tuple[int, int, str, str, float, str]] = []
    results.extend(_company_designator_results(text))
    results.extend(_contextual_organization_results(text))
    rules = (
        (
            "会社名",
            r"(?:会社名|法人名|取引先|顧客企業)[:：]\s*([一-龯々〆ヵヶぁ-んァ-ヶーA-Za-z0-9]{2,24})",
            "word_company_label",
            0.82,
        ),
        (
            "銀行名",
            r"(?:銀行名|金融機関|振込先|口座)[:：は\s]*([一-龯々〆ヵヶぁ-んァ-ヶーA-Za-z0-9]{2,20}(?:銀行|信用金庫|信用組合))",
            "word_bank_label",
            0.90,
        ),
        (
            "銀行名",
            r"([一-龯々〆ヵヶぁ-んァ-ヶーA-Za-z0-9]{2,20}(?:銀行|信用金庫|信用組合))",
            "word_bank_name",
            0.86,
        ),
        (
            "氏名",
            r"([一-龯々〆ヵヶ]{1,4}[ 　]+[一-龯々〆ヵヶ]{1,5})",
            "word_japanese_full_name_space",
            0.70,
        ),
        (
            "氏名",
            r"(?:担当者|連絡者|確認者|責任者|代表者)\s*[:：は]\s*([一-龯々〆ヵヶ]{2,6})(?![ 　])",
            "word_person_role_label",
            0.76,
        ),
    )
    for category, pattern, rule, confidence in rules:
        for match in re.finditer(pattern, text):
            start, end = match.span(1) if match.lastindex else match.span()
            start, end = _trim_candidate_span(text, start, end, category)
            value = text[start:end]
            if _is_generic_word_false_positive(category, value):
                continue
            results.append((start, end, category, rule, confidence, category))
    return results


def _company_designator_results(text: str) -> list[tuple[int, int, str, str, float, str]]:
    results: list[tuple[int, int, str, str, float, str]] = []
    designator_pattern = "|".join(re.escape(value) for value in COMPANY_DESIGNATORS)
    suffix_pattern = "|".join(re.escape(value) for value in COMPANY_SUFFIX_DESIGNATORS)

    for match in re.finditer(rf"(?:{designator_pattern})[{WORD_NAME_CHARS}]{{2,24}}", text):
        start, end = _trim_candidate_span(text, match.start(), match.end(), "会社名")
        value = text[start:end]
        if _has_company_name_body(value) and not _is_generic_word_false_positive("会社名", value):
            results.append((start, end, "会社名", "word_company_prefix", 0.88, "会社名"))

    for match in re.finditer(rf"[{WORD_NAME_CHARS}]{{2,30}}(?:{suffix_pattern})", text):
        suffix_start = _company_suffix_start(text, match.start(), match.end())
        start, end = _trim_candidate_span(text, suffix_start, match.end(), "会社名")
        value = text[start:end]
        if _has_company_name_body(value) and not _is_generic_word_false_positive("会社名", value):
            results.append((start, end, "会社名", "word_company_suffix", 0.86, "会社名"))
    return results


def _company_suffix_start(text: str, start: int, end: int) -> int:
    segment = text[start:end]
    adjusted = start
    for index, char in enumerate(segment):
        if char in " \t　:：、。，,/／「」『』（）()はをがにへでと":
            adjusted = start + index + 1
    return adjusted


def _has_company_name_body(value: str) -> bool:
    normalized = value.strip()
    for designator in COMPANY_DESIGNATORS:
        if normalized.startswith(designator):
            return len(normalized) - len(designator) >= 2
        if normalized.endswith(designator):
            return len(normalized) - len(designator) >= 2
    return False


def _contextual_organization_results(text: str) -> list[tuple[int, int, str, str, float, str]]:
    results: list[tuple[int, int, str, str, float, str]] = []
    context_pattern = r"(?:組織として|取引先(?:として|は)?|販売先(?:として|は)?|委託先(?:として|は)?|共同研究先(?:として|は)?)\s*"
    organization_pattern = rf"([{WORD_NAME_CHARS}]{{2,20}}(?:ラボ|研究所|センター|支店|本社|工場))"
    for match in re.finditer(context_pattern + organization_pattern, text):
        start, end = _trim_candidate_span(text, match.start(1), match.end(1), "会社名")
        value = text[start:end]
        if not any(designator in value for designator in COMPANY_DESIGNATORS) and not _is_generic_word_false_positive("会社名", value):
            results.append((start, end, "会社名", "word_contextual_organization", 0.72, "会社名"))
    return results


def _url_candidate_results(url: str) -> list[tuple[int, int, str, str, float]]:
    results: list[tuple[int, int, str, str, float]] = []
    for match in re.finditer(r"[A-Za-z0-9._%+\-]+@[A-Za-z0-9.\-]+\.[A-Za-z]{2,}", url):
        results.append((match.start(), match.end(), "メールアドレス", "url_email", 0.92))
    for match in re.finditer(r"(?:name|person|user|contact)=([A-Za-z][A-Za-z.\-_%]{2,60})", url, flags=re.IGNORECASE):
        start, end = match.span(1)
        results.append((start, end, "氏名", "url_name_like_parameter", 0.45))
    return results


def _select_non_overlapping_word_results(
    text: str,
    results: list[tuple[int, int, str, str, float, str]],
) -> list[tuple[int, int, str, str, float, str]]:
    candidates = [
        item
        for item in results
        if 0 <= item[0] < item[1] <= len(text) and len(text[item[0] : item[1]].strip()) >= 2
    ]
    pattern_based = [item for item in candidates if item[3] != WORD_NLP_DETECTION_RULE]
    nlp_based = [item for item in candidates if item[3] == WORD_NLP_DETECTION_RULE]

    pattern_based.sort(key=lambda item: (item[4], item[1] - item[0]), reverse=True)
    selected: list[tuple[int, int, str, str, float, str]] = []
    occupied_by_category: dict[str, list[range]] = {}
    for item in pattern_based:
        start, end, category, *_rest = item
        current = range(start, end)
        if any(current.start < existing.stop and existing.start < current.stop for existing in occupied_by_category.get(category, [])):
            continue
        selected.append(item)
        occupied_by_category.setdefault(category, []).append(current)

    # NLP-sourced candidates only fill spans no pattern/dictionary rule has
    # already claimed, in any category -- they exist to catch what those
    # rules structurally cannot see, not to duplicate an existing hit under a
    # different category label (e.g. a bank name GiNZA also tags as a
    # generic organization).
    occupied_any = [range(item[0], item[1]) for item in selected]
    nlp_based.sort(key=lambda item: item[1] - item[0], reverse=True)
    for item in nlp_based:
        start, end, *_rest = item
        current = range(start, end)
        if any(current.start < existing.stop and existing.start < current.stop for existing in occupied_any):
            continue
        selected.append(item)
        occupied_any.append(current)

    selected.sort(key=lambda item: (item[0], item[1], item[2]))
    return selected


def _normalize_word_category(category: str) -> str:
    mapping = {
        "PERSON": "氏名",
        "個人名": "氏名",
        "氏名候補": "氏名",
        "組織名": "会社名",
        "法人名": "会社名",
        "会社": "会社名",
        "電話": "電話番号",
        "メール": "メールアドレス",
        "銀行口座": "銀行名",
        "金融機関": "銀行名",
    }
    return mapping.get(category, category)


def _trim_candidate_span(text: str, start: int, end: int, category: str) -> tuple[int, int]:
    value = text[start:end]
    prefix_patterns = {
        "会社名": r"^(?:会社名|法人名|取引先|顧客企業)[:：は\s]*",
        "住所": r"^(?:送付先住所|住所|所在地|送付先|連絡先)[:：は\s]*",
        "銀行名": r"^(?:銀行名|金融機関|振込先|口座)[:：は\s]*",
    }
    prefix_match = re.match(prefix_patterns.get(category, r"$^"), value)
    if prefix_match:
        start += prefix_match.end()
        value = text[start:end]

    if category in {"会社名", "住所", "銀行名"}:
        stop_match = re.search(r"(?:です|でした|ます|ました|[。、「」,，]|(?<=[一-龯々〆ヵヶぁ-んァ-ヶーA-Za-z0-9])(?:は|を|が|に|へ|で|と)(?=[一-龯々〆ヵヶぁ-んァ-ヶー]))", value)
        if stop_match and stop_match.start() > 0:
            end = start + stop_match.start()
            value = text[start:end]

    if category in {"会社名", "氏名"}:
        for suffix in ("様", "さん", "殿", "氏"):
            if value.endswith(suffix) and len(value) > len(suffix):
                end -= len(suffix)
                value = text[start:end]
                break

    while start < end and text[start] in " \t　:：":
        start += 1
    while end > start and text[end - 1] in " \t　。、「」,，です":
        end -= 1
    return start, end


def _is_generic_word_false_positive(category: str, value: str) -> bool:
    normalized = value.strip()
    if category == "会社名" and normalized in set(COMPANY_DESIGNATORS):
        return True
    if category == "銀行名" and normalized in {"銀行", "銀行名", "金融機関"}:
        return True
    if category == "会社名" and normalized in {"会社", "会社名", "法人名", "担当", "担当者"}:
        return True
    if category == "会社名" and not any(word in normalized for word in COMPANY_DESIGNATORS) and re.search(r"(?:担当者?|先)$", normalized) and len(normalized) <= 8:
        return True
    if category == "氏名" and any(word in normalized for word in ("銀行", "支店", "会社", "法人", "研究所", "センター", "ラボ", "工場")):
        return True
    return False


def _extend_address_span(text: str, start: int, end: int) -> tuple[int, int]:
    if not _address_has_strength(text[start:end]):
        return start, end
    suffix = text[end:]
    building_match = re.match(
        r"[ \t　]+(?:"
        r"[一-龯々〆ヵヶぁ-んァ-ヶーA-Za-z0-9０-９\-－]{1,20}"
        r"(?:ビル|マンション|号館|棟|タワー|ハイツ|レジデンス)"
        r"[一-龯々〆ヵヶぁ-んァ-ヶーA-Za-z0-9０-９\-－]*(?:[0-9０-９一二三四五六七八九十百千]+(?:階|F|Ｆ|号室))?"
        r"|[0-9０-９一二三四五六七八九十百千]+(?:階|F|Ｆ|号室)"
        r")",
        suffix,
    )
    if building_match:
        end += building_match.end()
    return start, end


def _is_weak_address_candidate(value: str) -> bool:
    normalized = value.strip()
    if not _address_has_strength(normalized):
        return True
    if re.search(r"(?:制度|政策|説明|について)", normalized):
        return True
    return False


def _address_has_strength(value: str) -> bool:
    return bool(
        re.search(r"[0-9０-９]", value)
        or re.search(r"(?:丁目|番地|番|号|号室|〒)", value)
        or re.search(r"[一二三四五六七八九十百千]+(?:丁目|番地|番|号)", value)
    )


def _affected_run_indices(runs: tuple[WordTextRun, ...], start: int, end: int) -> tuple[int, ...]:
    return tuple(
        run.run_index
        for run in runs
        if start < run.char_end and run.char_start < end
    )


def _runs_by_paragraph_location(runs: tuple[WordTextRun, ...]) -> dict[str, tuple[WordTextRun, ...]]:
    grouped: dict[str, list[WordTextRun]] = {}
    for run in runs:
        grouped.setdefault(_run_paragraph_location_id(run), []).append(run)
    return {key: tuple(sorted(value, key=lambda item: item.run_index)) for key, value in grouped.items()}


def _paragraph_location_id(paragraph: WordParagraphText) -> str:
    values = [
        paragraph.story_type,
        str(paragraph.section_index),
        paragraph.header_footer_type or "",
        paragraph.container_type,
        str(paragraph.table_index),
        str(paragraph.row_index),
        str(paragraph.cell_index),
        str(paragraph.paragraph_index),
        str(paragraph.cell_paragraph_index),
        paragraph.part_name,
        paragraph.element_path,
    ]
    return ":".join(values)


def _run_paragraph_location_id(run: WordTextRun) -> str:
    values = [
        run.story_type,
        str(run.section_index),
        run.header_footer_type or "",
        run.container_type,
        str(run.table_index),
        str(run.row_index),
        str(run.cell_index),
        str(run.paragraph_index),
        str(run.cell_paragraph_index),
        run.part_name,
        run.element_path.rsplit("/r[", 1)[0],
    ]
    return ":".join(values)


def _candidate_id(
    source_sha256: str,
    location_id: str,
    category: str,
    start: int,
    end: int,
    text: str,
    source: str,
) -> str:
    payload = "|".join([source_sha256, location_id, category, str(start), str(end), text, source])
    return hashlib.sha256(payload.encode("utf-8")).hexdigest()


def _document_property_values(properties: WordDocumentProperties) -> tuple[tuple[str, str], ...]:
    return (
        ("author", properties.author),
        ("last_modified_by", properties.last_modified_by),
        ("title", properties.title),
        ("subject", properties.subject),
        ("keywords", properties.keywords),
        ("category", properties.category),
        ("comments", properties.comments),
        ("content_status", properties.content_status),
        ("identifier", properties.identifier),
        ("language", properties.language),
        ("version", properties.version),
    )


def _inspect_xml_text(
    part_name: str,
    text: str,
    features: dict[tuple[str, str], int],
    details: dict[tuple[str, str], str],
) -> None:
    checks = (
        ("textbox", ("txbxContent", "v:textbox", "wps:txbx")),
        ("drawing_or_shape", ("<w:drawing", "<v:shape", "<wps:wsp")),
        ("tracked_changes", ("<w:ins", "<w:del", "<w:moveFrom", "<w:moveTo")),
        ("complex_field_code", ("<w:fldChar", "<w:instrText")),
    )
    for feature_type, markers in checks:
        count = sum(text.count(marker) for marker in markers)
        if count:
            _count_feature(features, details, feature_type, part_name, count=count)
    if part_name.endswith(".rels") and 'TargetMode="External"' in text:
        _count_feature(features, details, "external_link", part_name, count=text.count('TargetMode="External"'))
    if part_name.endswith(".xml") and _may_hold_text(part_name, text):
        _count_feature(features, details, "text_holding_xml_part", part_name)


def _may_hold_text(part_name: str, text: str) -> bool:
    supported_prefixes = (
        "word/document.xml",
        "word/header",
        "word/footer",
        "docProps/",
    )
    if part_name.startswith(supported_prefixes):
        return False
    return "<w:t" in text or "<dc:" in text or "<cp:" in text


def _count_feature(
    features: dict[tuple[str, str], int],
    details: dict[tuple[str, str], str],
    feature_type: str,
    part_name: str,
    *,
    count: int = 1,
    detail: str = "",
) -> None:
    key = (feature_type, part_name)
    features[key] = features.get(key, 0) + count
    if detail:
        details[key] = detail


def _part_name(item: Any) -> str:
    part = getattr(item, "part", None)
    if part is None:
        parent = getattr(item, "_parent", None)
        part = getattr(parent, "part", None)
    partname = getattr(part, "partname", "")
    return str(partname)


def _local_name(tag: str) -> str:
    return tag.rsplit("}", 1)[-1] if "}" in tag else tag


def _file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def _hmac_digest(value: str) -> str:
    return hmac.new(_state_hmac_key(), value.encode("utf-8"), hashlib.sha256).hexdigest()


def _state_hmac_key() -> bytes:
    env_key = os.environ.get("PRIVACY_CLEANER_STATE_HMAC_KEY")
    if env_key:
        return env_key.encode("utf-8")

    key_path = _state_hmac_key_path()
    if key_path.exists():
        return key_path.read_bytes()
    key_path.parent.mkdir(parents=True, exist_ok=True)
    key = secrets.token_bytes(32)
    key_path.write_bytes(key)
    return key


def _state_hmac_key_path() -> Path:
    base = Path(os.environ.get("LOCALAPPDATA", Path.cwd()))
    preferred = base / "ExcelPrivacyCleaner" / "state_hmac.key"
    try:
        preferred.parent.mkdir(parents=True, exist_ok=True)
        probe = preferred.parent / ".write_test"
        probe.write_text("", encoding="utf-8")
        probe.unlink(missing_ok=True)
        return preferred
    except Exception:
        return Path.cwd() / "config" / ".privacy_cleaner_state_hmac.key"


def _word_hmac_digest(value: str) -> str:
    return hmac.new(_state_hmac_key(), f"{WORD_AUDIT_SCHEMA}:{value}".encode("utf-8"), hashlib.sha256).hexdigest()
