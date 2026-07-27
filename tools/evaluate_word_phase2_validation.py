from __future__ import annotations

import csv
import hashlib
import json
import sys
import tempfile
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from excel_privacy_cleaner.word_processor import WordCandidate, WordParagraphText, candidates_for_inventory, extract_word_structure  # noqa: E402


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs" / "evaluation_baselines" / "word_phase2_validation"
VALIDATION_DOCX = OUTPUT_DIR / "Word匿名化検出_Phase2_ValidationSynthetic.docx"
GROUND_TRUTH_CSV = OUTPUT_DIR / "Word匿名化検出_Phase2_ValidationGroundTruth.csv"
INITIAL_CANDIDATES_CSV = OUTPUT_DIR / "Word匿名化検出_Phase2_ValidationInitialCandidates.csv"
EVALUATION_JSON = ROOT / "docs" / "evaluation_baselines" / "Word匿名化検出_Phase2_ValidationInitial.json"
EVALUATION_CSV = ROOT / "docs" / "evaluation_baselines" / "Word匿名化検出_Phase2_ValidationInitial.csv"
EVALUATION_MD = ROOT / "docs" / "evaluation_baselines" / "Word匿名化検出_Phase2_ValidationInitial.md"
DATASET_CONFIG = ROOT / "config" / "evaluation" / "word_dataset_split_v1.json"


@dataclass(frozen=True)
class TruthSpec:
    category: str
    marker: str
    occurrence: int = 0
    source: str = "paragraph"


VALIDATION_TRUTHS = (
    TruthSpec("会社名", "株式会社星雲商事", 0),
    TruthSpec("会社名", "北斗物流株式会社", 0),
    TruthSpec("会社名", "合同会社青葉研究所", 0),
    TruthSpec("会社名", "白樺ラボ", 0),
    TruthSpec("氏名", "田中 一郎", 0),
    TruthSpec("氏名", "鈴木花子", 0),
    TruthSpec("氏名", "高橋 次郎", 0),
    TruthSpec("氏名", "佐々木 三郎", 0),
    TruthSpec("住所", "北海道札幌市中央区北一条西2-3", 0),
    TruthSpec("住所", "名古屋市中区錦3丁目4-5", 0),
    TruthSpec("住所", "京都府京都市下京区烏丸通1-2 青葉ビル5階", 0),
    TruthSpec("電話番号", "03-1234-5678", 0),
    TruthSpec("電話番号", "09012345678", 0),
    TruthSpec("メールアドレス", "info@seiun.example.jp", 0),
    TruthSpec("メールアドレス", "support@sub.aoba.example.jp", 0),
    TruthSpec("メールアドレス", "long.contact-address+phase2@example.co.jp", 0),
    TruthSpec("銀行名", "北都銀行", 0),
    TruthSpec("銀行名", "青葉信用金庫", 0),
    TruthSpec("銀行名", "中央信用組合", 0),
    TruthSpec("会社名", "株式会社星雲商事", 1),
    TruthSpec("氏名", "田中 一郎", 1),
    TruthSpec("会社名", "北斗物流株式会社", 1),
    TruthSpec("氏名", "佐々木 三郎", 1),
    TruthSpec("会社名", "株式会社星雲商事", 2),
)


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def create_validation_docx(path: Path) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    document = Document()
    document.core_properties.author = "Validation Synthetic"
    document.core_properties.title = "Word Phase 2 Validation Synthetic"
    document.add_paragraph("取引先は株式会社星雲商事です。")
    document.add_paragraph("配送担当は北斗物流株式会社です。")
    document.add_paragraph("共同研究先は合同会社青葉研究所です。")
    document.add_paragraph("法人格を含まない組織として白樺ラボを記載します。")

    split_company = document.add_paragraph()
    split_company.add_run("株式会社")
    split_company.add_run("星雲")
    split_company.add_run("商事")
    split_company.add_run(" は複数runで再登場します。")

    document.add_paragraph("担当者は田中 一郎です。")
    document.add_paragraph("確認済み担当者として田中 一郎が再登場します。")
    document.add_paragraph("連絡者は鈴木花子です。")
    split_name = document.add_paragraph()
    split_name.add_run("確認者は")
    split_name.add_run("高橋")
    split_name.add_run(" 次郎")
    split_name.add_run("様です。")
    document.add_paragraph("一般語として南側の入口、森を抜ける、銀行の役割、会社の方針を記載します。")

    document.add_paragraph("住所1: 北海道札幌市中央区北一条西2-3")
    document.add_paragraph("住所2: 名古屋市中区錦3丁目4-5")
    document.add_paragraph("住所3: 京都府京都市下京区烏丸通1-2 青葉ビル5階")
    document.add_paragraph("住所に似た一般文として、市区町村の制度について説明します。")

    document.add_paragraph("固定電話 03-1234-5678 / 携帯 09012345678")
    document.add_paragraph("対象外数字 123456789 と 03-12-345 は管理番号です。")
    document.add_paragraph("メール info@seiun.example.jp / support@sub.aoba.example.jp / long.contact-address+phase2@example.co.jp")
    document.add_paragraph("不正メール user@@example.jp と missing-at.example.jp は対象外です。")
    document.add_paragraph("振込先は北都銀行、青葉信用金庫、中央信用組合です。銀行という一般語は対象外です。")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "会社"
    table.cell(0, 1).text = "氏名"
    table.cell(1, 0).text = "北斗物流株式会社"
    table.cell(1, 1).text = "佐々木 三郎"

    section = document.sections[0]
    section.header.paragraphs[0].text = "ヘッダー 株式会社星雲商事"
    section.footer.paragraphs[0].text = "フッター 佐々木 三郎"
    document.save(path)


def main() -> int:
    create_validation_docx(VALIDATION_DOCX)
    inventory = extract_word_structure(VALIDATION_DOCX)
    truths = build_truth_from_inventory(inventory.paragraphs, VALIDATION_TRUTHS)
    write_ground_truth(GROUND_TRUTH_CSV, truths)
    ground_truth_sha256 = file_sha256(GROUND_TRUTH_CSV)

    candidates = candidates_for_inventory(inventory)
    write_candidates(INITIAL_CANDIDATES_CSV, candidates)
    metrics, category_results = evaluate(candidates, truths)
    candidate_sha256 = file_sha256(INITIAL_CANDIDATES_CSV)
    validation_docx_sha256 = file_sha256(VALIDATION_DOCX)
    update_dataset_config(validation_docx_sha256, ground_truth_sha256)
    write_evaluation_records(metrics, category_results, validation_docx_sha256, ground_truth_sha256, candidate_sha256)

    print("word_phase2_validation_initial=completed")
    print("validation_docx_sha256", validation_docx_sha256)
    print("ground_truth_sha256", ground_truth_sha256)
    print("candidate_sha256", candidate_sha256)
    print("metrics", json.dumps(metrics, ensure_ascii=False, sort_keys=True))
    print("category_results", json.dumps(category_results, ensure_ascii=False, sort_keys=True))
    return 0


def build_truth_from_inventory(paragraphs: Iterable[WordParagraphText], truth_specs: tuple[TruthSpec, ...]) -> list[dict[str, object]]:
    truths: list[dict[str, object]] = []
    for index, spec in enumerate(truth_specs, start=1):
        matches: list[tuple[WordParagraphText, int, int]] = []
        for paragraph in paragraphs:
            if spec.source != "paragraph" or paragraph.story_type == "document_property":
                continue
            start = 0
            while True:
                found = paragraph.text.find(spec.marker, start)
                if found < 0:
                    break
                matches.append((paragraph, found, found + len(spec.marker)))
                start = found + 1
        assert_true(len(matches) > spec.occurrence, f"Missing validation truth marker: {spec}")
        paragraph, start, end = matches[spec.occurrence]
        truths.append(
            {
                "truth_id": f"WORDVAL-{index:03d}",
                "category": spec.category,
                "location_id": paragraph_location_id(paragraph),
                "char_start": start,
                "char_end": end,
                "source": spec.source,
            }
        )
    return truths


def evaluate(candidates: tuple[WordCandidate, ...], truths: list[dict[str, object]]) -> tuple[dict[str, object], dict[str, dict[str, int]]]:
    truth_keys = {
        (str(truth["location_id"]), str(truth["category"]), int(truth["char_start"]), int(truth["char_end"]))
        for truth in truths
    }
    candidate_keys = {
        (candidate.location_id, candidate.category, candidate.char_start, candidate.char_end)
        for candidate in candidates
    }
    matched = truth_keys & candidate_keys
    missed = truth_keys - candidate_keys
    true_positive = candidate_keys & truth_keys
    false_positive = candidate_keys - truth_keys
    metrics = {
        "ground_truth_count": len(truth_keys),
        "candidate_count": len(candidate_keys),
        "matched_ground_truth_count": len(matched),
        "missed_ground_truth_count": len(missed),
        "true_positive_candidate_count": len(true_positive),
        "false_positive_candidate_count": len(false_positive),
        "recall": round(len(matched) / len(truth_keys), 3) if truth_keys else 0.0,
        "candidate_precision": round(len(true_positive) / len(candidate_keys), 3) if candidate_keys else 0.0,
        "exact_span_match_count": len(matched),
    }
    category_results: dict[str, dict[str, int]] = defaultdict(lambda: {"正解": 0, "検出": 0, "見逃し": 0, "TP候補": 0, "FP候補": 0})
    for _location, category, _start, _end in truth_keys:
        category_results[category]["正解"] += 1
    for _location, category, _start, _end in candidate_keys:
        category_results[category]["検出"] += 1
    for _location, category, _start, _end in missed:
        category_results[category]["見逃し"] += 1
    for _location, category, _start, _end in true_positive:
        category_results[category]["TP候補"] += 1
    for _location, category, _start, _end in false_positive:
        category_results[category]["FP候補"] += 1
    return metrics, dict(sorted(category_results.items()))


def write_ground_truth(path: Path, truths: list[dict[str, object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as output:
        writer = csv.DictWriter(output, fieldnames=["truth_id", "category", "location_id", "char_start", "char_end", "source"])
        writer.writeheader()
        for truth in truths:
            writer.writerow(truth)


def write_candidates(path: Path, candidates: tuple[WordCandidate, ...]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as output:
        writer = csv.DictWriter(
            output,
            fieldnames=["candidate_id", "category", "location_id", "char_start", "char_end", "source", "detection_rule", "confidence"],
        )
        writer.writeheader()
        for candidate in candidates:
            writer.writerow(
                {
                    "candidate_id": candidate.candidate_id,
                    "category": candidate.category,
                    "location_id": candidate.location_id,
                    "char_start": candidate.char_start,
                    "char_end": candidate.char_end,
                    "source": candidate.source,
                    "detection_rule": candidate.detection_rule,
                    "confidence": f"{candidate.confidence:.3f}",
                }
            )


def write_evaluation_records(
    metrics: dict[str, object],
    category_results: dict[str, dict[str, int]],
    validation_docx_sha256: str,
    ground_truth_sha256: str,
    candidate_sha256: str,
) -> None:
    payload = {
        "baseline_name": "Word candidate detection Phase 2 validation initial",
        "source_commit": "441d588",
        "validation_docx_sha256": validation_docx_sha256,
        "ground_truth_sha256": ground_truth_sha256,
        "initial_candidates_sha256": candidate_sha256,
        "metrics": metrics,
        "category_results": category_results,
        "method": {
            "formal_true_positive": "same location_id, normalized category, exact char_start, exact char_end",
            "rule_changes_after_viewing_validation": "none",
        },
        "limitations": [
            "Synthetic validation data only.",
            "No rule changes were made after fixing validation ground truth.",
            "This is not a holdout evaluation."
        ],
    }
    EVALUATION_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")
    with EVALUATION_CSV.open("w", encoding="utf-8-sig", newline="") as output:
        writer = csv.writer(output)
        writer.writerow(["項目", "値", "備考"])
        for key, value in metrics.items():
            writer.writerow([key, value, ""])
        writer.writerow(["validation_docx_sha256", validation_docx_sha256, ""])
        writer.writerow(["ground_truth_sha256", ground_truth_sha256, ""])
        writer.writerow(["initial_candidates_sha256", candidate_sha256, ""])
        for category, result in category_results.items():
            writer.writerow([f"category_{category}", json.dumps(result, ensure_ascii=False), ""])
    EVALUATION_MD.write_text(_evaluation_markdown(metrics, category_results, validation_docx_sha256, ground_truth_sha256), encoding="utf-8")


def _evaluation_markdown(
    metrics: dict[str, object],
    category_results: dict[str, dict[str, int]],
    validation_docx_sha256: str,
    ground_truth_sha256: str,
) -> str:
    lines = [
        "# Word匿名化検出 Phase2 Validation Initial",
        "",
        "Phase 2 baselineコードを変更せず、validation合成Wordに対して初回候補生成を実行した記録です。",
        "",
        f"- validation_docx_sha256: `{validation_docx_sha256}`",
        f"- ground_truth_sha256: `{ground_truth_sha256}`",
        "- 評価方法: 同一location_id、カテゴリ一致、char_start完全一致、char_end完全一致を正式TPとする。",
        "- この結果を見たルール改善はまだ実施していません。",
        "",
        "## Metrics",
        "",
        "| 指標 | 値 |",
        "|---|---:|",
    ]
    for key, value in metrics.items():
        lines.append(f"| {key} | `{value}` |")
    lines.extend(["", "## カテゴリ別結果", "", "| カテゴリ | 正解 | 検出 | 見逃し | TP候補 | FP候補 |", "|---|---:|---:|---:|---:|---:|"])
    for category, result in category_results.items():
        lines.append(
            f"| {category} | `{result['正解']}` | `{result['検出']}` | `{result['見逃し']}` | `{result['TP候補']}` | `{result['FP候補']}` |"
        )
    lines.append("")
    return "\n".join(lines)


def update_dataset_config(validation_docx_sha256: str, ground_truth_sha256: str) -> None:
    payload = json.loads(DATASET_CONFIG.read_text(encoding="utf-8"))
    payload["validation"] = {
        "synthetic_phase2_docx": {
            "path": str(VALIDATION_DOCX.relative_to(ROOT)),
            "sha256": validation_docx_sha256,
        },
        "ground_truth_csv": {
            "path": str(GROUND_TRUTH_CSV.relative_to(ROOT)),
            "sha256": ground_truth_sha256,
        },
        "note": "Validation ground truth is fixed before initial candidate generation. Do not use for rule tuning until the initial result is recorded.",
    }
    DATASET_CONFIG.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def paragraph_location_id(paragraph: WordParagraphText) -> str:
    values = [
        paragraph.story_type,
        str(paragraph.section_index),
        paragraph.header_footer_type or "",
        paragraph.container_type,
        str(paragraph.table_index),
        str(paragraph.row_index),
        str(paragraph.cell_index),
        str(paragraph.paragraph_index),
        str(paragraph.cell_paragraph_index),
        paragraph.part_name,
        paragraph.element_path,
    ]
    return ":".join(values)


def file_sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


if __name__ == "__main__":
    raise SystemExit(main())
