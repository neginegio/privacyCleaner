from __future__ import annotations

import re
import sys
import tempfile
import zipfile
from dataclasses import replace
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.shared import Pt, RGBColor

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from excel_privacy_cleaner.excel_processor import AliasBook, ProcessingOptions  # noqa: E402
from excel_privacy_cleaner.word_processor import (  # noqa: E402
    WordCandidate,
    WordPrivacyProcessor,
    WordReplacementDecision,
    build_default_decisions,
    candidate_requires_review,
    default_replacement_for_candidate,
    extract_word_structure,
    _find_residual_text,
)


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:val"), "clear")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def create_replacement_fixture(path: Path) -> None:
    document = Document()
    document.core_properties.author = "山田 太郎"
    document.core_properties.last_modified_by = "株式会社未来会議"
    document.core_properties.title = "置換テスト用文書"

    phone_paragraph = document.add_paragraph()
    phone_paragraph.add_run("090-1111-2222")
    note_run = phone_paragraph.add_run("(至急)")
    note_run.bold = True
    note_run.font.size = Pt(14)
    note_run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

    mixed_paragraph = document.add_paragraph()
    mixed_paragraph.add_run("株式会社")
    mixed_paragraph.add_run("未来会議")
    mixed_paragraph.add_run(" の担当は")
    mixed_paragraph.add_run("山田")
    mixed_paragraph.add_run(" 太郎")
    mixed_paragraph.add_run("です。")

    document.add_paragraph("振込先はみらい銀行 東京支店です。")

    linked_paragraph = document.add_paragraph("リンク先も確認: ")
    add_hyperlink(linked_paragraph, "会社サイト", "https://example.invalid/contact?email=contact@example.jp")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "会社"
    table.cell(0, 1).text = "氏名"
    table.cell(1, 0).text = "株式会社未来会議"
    add_cell_shading(table.cell(1, 0), "FFEEAA")
    table.cell(1, 1).text = "佐藤 花子"

    section = document.sections[0]
    section.header.paragraphs[0].text = "ヘッダー 株式会社未来会議"
    section.footer.paragraphs[0].text = "フッター 山田 太郎"
    document.save(path)


def create_low_confidence_name_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("佐藤 花子")
    document.save(path)


def create_high_confidence_only_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("連絡先電話は090-1111-2222です。")
    document.save(path)


def inject_unsupported_comments_part(path: Path) -> None:
    tmp_path = path.with_name(path.name + ".tmp")
    with zipfile.ZipFile(path) as source_archive, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as target_archive:
        for info in source_archive.infolist():
            target_archive.writestr(info, source_archive.read(info.filename))
        target_archive.writestr(
            "word/comments.xml",
            (
                '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                '<w:comments xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
                '<w:comment w:id="0"><w:p><w:r><w:t>コメント内文字</w:t></w:r></w:p></w:comment>'
                "</w:comments>"
            ).encode("utf-8"),
        )
    tmp_path.replace(path)


def inject_company_into_app_properties(path: Path, company_text: str) -> None:
    tmp_path = path.with_name(path.name + ".tmp")
    with zipfile.ZipFile(path) as source_archive, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as target_archive:
        for info in source_archive.infolist():
            data = source_archive.read(info.filename)
            if info.filename == "docProps/app.xml":
                text = data.decode("utf-8")
                text = re.sub(r"<Company\s*/>|<Company>.*?</Company>", f"<Company>{company_text}</Company>", text)
                data = text.encode("utf-8")
            target_archive.writestr(info, data)
    tmp_path.replace(path)


def _decisions_by_text(decisions: list[WordReplacementDecision], text: str) -> list[WordReplacementDecision]:
    return [decision for decision in decisions if decision.candidate.text == text]


def _enable_review_required(decisions: list[WordReplacementDecision]) -> list[WordReplacementDecision]:
    # The fixture intentionally includes low-confidence (<0.75) name candidates
    # (table cell, footer, author property) that default to disabled under the
    # review-required guard. Tests that aren't about that guard itself approve
    # them explicitly here, mirroring a reviewer confirming a real candidate.
    for decision in decisions:
        if decision.candidate.source != "hyperlink_target":
            decision.enabled = True
    return decisions


def test_single_run_replacement_preserves_sibling_run_formatting() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        before = Document(source)
        note_run_before = before.paragraphs[0].runs[1]
        before_format = (note_run_before.bold, note_run_before.font.size, note_run_before.font.color.rgb, note_run_before.text)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        result = processor.convert(source, decisions, output_dir=tmp)

        assert_true("090-1111-2222" not in Document(result.output_path).paragraphs[0].runs[0].text, "Phone run should be replaced")

        after = Document(result.output_path)
        note_run_after = after.paragraphs[0].runs[1]
        after_format = (note_run_after.bold, note_run_after.font.size, note_run_after.font.color.rgb, note_run_after.text)
        assert_true(before_format == after_format, "Sibling run formatting and text must be untouched")


def test_multi_run_candidate_replaces_first_run_and_empties_rest() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        result = processor.convert(source, decisions, output_dir=tmp)

        output_paragraphs = Document(result.output_path).paragraphs
        mixed = output_paragraphs[1]
        assert_true(len(mixed.runs) == 6, "Run count in the mixed paragraph must not change")
        assert_true(mixed.runs[0].text.startswith("法人"), "First run of split company should hold the alias")
        assert_true(mixed.runs[1].text == "", "Second run of split company should be emptied, not deleted")
        assert_true("株式会社未来会議" not in mixed.text, "Original company text should not remain")
        assert_true("山田" not in mixed.text and "太郎" not in mixed.text, "Original name text should not remain")


def test_two_non_overlapping_candidates_in_same_paragraph() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        def in_mixed_paragraph(decision: WordReplacementDecision) -> bool:
            candidate = decision.candidate
            return candidate.source == "paragraph" and candidate.story_type == "body" and candidate.container_type == "paragraph"

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        company_decisions = [d for d in _decisions_by_text(decisions, "株式会社未来会議") if in_mixed_paragraph(d)]
        name_decisions = [d for d in _decisions_by_text(decisions, "山田 太郎") if in_mixed_paragraph(d)]
        assert_true(len(company_decisions) == 1 and len(name_decisions) == 1, "Both candidates should be detected once in this paragraph")

        result = processor.convert(source, decisions, output_dir=tmp)
        mixed_text = Document(result.output_path).paragraphs[1].text
        assert_true(mixed_text.startswith("法人"), "Company alias should be present")
        assert_true("個人" in mixed_text, "Name alias should be present")
        assert_true("の担当は" in mixed_text and "です。" in mixed_text, "Untouched surrounding text should remain")


def test_table_cell_replacement_preserves_cell_formatting() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        before_tc_pr = Document(source).tables[0].cell(1, 0)._tc.tcPr
        before_xml = before_tc_pr.xml if before_tc_pr is not None else None

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        result = processor.convert(source, decisions, output_dir=tmp)

        output_table = Document(result.output_path).tables[0]
        assert_true("株式会社未来会議" not in output_table.cell(1, 0).text, "Cell text should be replaced")
        assert_true("佐藤" not in output_table.cell(1, 1).text, "Second cell name should be replaced")
        after_tc_pr = output_table.cell(1, 0)._tc.tcPr
        after_xml = after_tc_pr.xml if after_tc_pr is not None else None
        assert_true(before_xml == after_xml, "Cell shading/borders must be byte-identical after replacement")


def test_header_and_footer_replacement() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        result = processor.convert(source, decisions, output_dir=tmp)

        output_document = Document(result.output_path)
        header_text = output_document.sections[0].header.paragraphs[0].text
        footer_text = output_document.sections[0].footer.paragraphs[0].text
        assert_true("株式会社未来会議" not in header_text and "ヘッダー" in header_text, "Header company should be replaced, prefix kept")
        assert_true("山田" not in footer_text and "フッター" in footer_text, "Footer name should be replaced, prefix kept")


def test_document_property_replacement() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        result = processor.convert(source, decisions, output_dir=tmp)

        properties = Document(result.output_path).core_properties
        assert_true("山田" not in properties.author, "Author property should be anonymized")
        assert_true("株式会社未来会議" not in properties.last_modified_by, "Last-modified-by property should be anonymized")
        assert_true(properties.title == "置換テスト用文書", "Title without PII should be untouched")


def test_alias_reused_for_repeated_original_text() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        company_aliases = {decision.replacement for decision in _decisions_by_text(decisions, "株式会社未来会議")}
        assert_true(len(company_aliases) == 1, "Same original text must map to a single alias within one scan")

        result = processor.convert(source, decisions, output_dir=tmp)
        output_document = Document(result.output_path)
        alias = next(iter(company_aliases))
        header_text = output_document.sections[0].header.paragraphs[0].text
        cell_text = output_document.tables[0].cell(1, 0).text
        assert_true(alias in header_text, "Alias should reach the header occurrence")
        assert_true(alias in cell_text, "Alias should reach the table-cell occurrence")


def test_bank_name_uses_financial_institution_alias_prefix() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        bank_decisions = _decisions_by_text(decisions, "みらい銀行")
        assert_true(len(bank_decisions) == 1, "Bank name should be detected once")
        assert_true(bank_decisions[0].replacement.startswith("金融機関"), "Bank alias should use the financial_institution prefix")

        result = processor.convert(source, decisions, output_dir=tmp)
        output_text = Document(result.output_path).paragraphs[2].text
        assert_true("みらい銀行" not in output_text and "金融機関" in output_text, "Bank name should be replaced in output")


def test_hyperlink_target_candidates_are_disabled_by_default_and_guarded() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        hyperlink_decisions = [d for d in decisions if d.candidate.source == "hyperlink_target"]
        assert_true(bool(hyperlink_decisions), "Hyperlink URL should produce at least one candidate")
        assert_true(all(not d.enabled for d in hyperlink_decisions), "Hyperlink URL candidates must default to disabled")

        _enable_review_required(decisions)
        result = processor.convert(source, decisions, output_dir=tmp)
        output_inventory = extract_word_structure(result.output_path)
        assert_true(
            any(link.url == "https://example.invalid/contact?email=contact@example.jp" for link in output_inventory.hyperlinks),
            "Disabled hyperlink URL candidate should leave the URL unchanged",
        )

        processor_2 = WordPrivacyProcessor()
        decisions_2 = processor_2.scan(source)
        _enable_review_required(decisions_2)
        for decision in decisions_2:
            if decision.candidate.source == "hyperlink_target":
                decision.enabled = True
        raised = False
        try:
            processor_2.convert(source, decisions_2, output_dir=tmp)
        except RuntimeError as exc:
            raised = True
            assert_true("contact@example.jp" in str(exc) or "ハイパーリンク" in str(exc), "Error should reference the guarded candidate")
        assert_true(raised, "Enabling a hyperlink_target candidate must raise, not silently apply")


def test_overlap_guard_rejects_overlapping_ranges() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        company_decision = _decisions_by_text(decisions, "株式会社未来会議")[0]
        candidate = company_decision.candidate
        overlapping_candidate = replace(
            candidate,
            candidate_id=candidate.candidate_id + "_synthetic_overlap",
            category="氏名",
            char_start=candidate.char_start + 2,
            char_end=candidate.char_end,
            text=candidate.text[2:],
            affected_run_indices=candidate.affected_run_indices,
        )
        overlapping_decision = WordReplacementDecision(candidate=overlapping_candidate, enabled=True, replacement="重複")

        raised = False
        try:
            processor.convert(source, decisions + [overlapping_decision], output_dir=tmp)
        except RuntimeError:
            raised = True
        assert_true(raised, "Overlapping candidate ranges in the same location must raise")


def test_integrity_guard_rejects_stale_candidate_text() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        company_decision = _decisions_by_text(decisions, "株式会社未来会議")[0]
        stale_candidate = replace(company_decision.candidate, text="ズレた文字列")
        stale_decisions = [d for d in decisions if d is not company_decision]
        stale_decisions.append(WordReplacementDecision(candidate=stale_candidate, enabled=True, replacement="法人999"))

        raised = False
        try:
            processor.convert(source, stale_decisions, output_dir=tmp)
        except RuntimeError:
            raised = True
        assert_true(raised, "A candidate whose text no longer matches the live paragraph must raise")


def test_end_to_end_scan_and_convert_leaves_no_marker_strings() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        result = processor.convert(source, decisions, output_dir=tmp)

        # This checks the readable body/table/header/footer text via python-docx;
        # the raw internal-XML residual scan (step 7) is asserted separately below.
        output_inventory = extract_word_structure(result.output_path)
        combined_text = "\n".join(paragraph.text for paragraph in output_inventory.paragraphs)
        for marker in ("株式会社未来会議", "山田", "太郎", "佐藤", "花子", "みらい銀行", "090-1111-2222"):
            assert_true(marker not in combined_text, f"Marker should not remain in output body/table/header/footer: {marker}")

        assert_true(
            not any("残存" in warning for warning in result.warnings),
            "Well-behaved fixture must not trigger the internal-XML residual warning",
        )


def test_default_replacement_for_candidate_uses_alias_book_directly() -> None:
    alias_book = AliasBook()
    options = ProcessingOptions()
    candidate = WordCandidate(
        candidate_id="synthetic",
        category="銀行名",
        story_type="body",
        section_index=None,
        container_type="paragraph",
        paragraph_index=0,
        table_index=None,
        row_index=None,
        cell_index=None,
        location_id="body::paragraph:None:None:None:0:None::",
        char_start=0,
        char_end=4,
        text="みらい銀行",
        detection_rule="test",
        confidence=0.9,
        source="paragraph",
    )
    replacement = default_replacement_for_candidate(candidate, alias_book, options)
    assert_true(replacement.startswith("金融機関"), "Bank category should route through the financial_institution alias kind")

    decisions = build_default_decisions((candidate,), alias_book, options)
    assert_true(decisions[0].replacement == replacement, "build_default_decisions should reuse the same alias")


def test_low_confidence_candidate_blocks_by_default() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_low_confidence_name_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        assert_true(len(decisions) == 1, "Fixture should yield exactly one candidate")
        name_decision = decisions[0]
        assert_true(name_decision.candidate.confidence < 0.75, "Unlabeled two-part name should be a low-confidence match")
        assert_true(not name_decision.enabled, "Low-confidence candidate must default to disabled")

        raised = False
        try:
            processor.convert(source, decisions, output_dir=tmp)
        except RuntimeError as exc:
            raised = True
            assert_true("未確認候補" in str(exc), "Error should identify the block as an unreviewed candidate")
        assert_true(raised, "Leaving a low-confidence candidate disabled must block conversion")


def test_review_required_candidate_allows_conversion_when_enabled() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_low_confidence_name_fixture(source)

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        decisions[0].enabled = True
        result = processor.convert(source, decisions, output_dir=tmp)

        output_text = Document(result.output_path).paragraphs[0].text
        assert_true("佐藤" not in output_text and "花子" not in output_text, "Approved candidate should be replaced")


def test_unsupported_feature_blocks_external_mode_but_warns_in_analysis_mode() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_high_confidence_only_fixture(source)
        inject_unsupported_comments_part(source)

        external_processor = WordPrivacyProcessor()
        external_decisions = external_processor.scan(source, options=ProcessingOptions(mode="external"))
        raised = False
        try:
            external_processor.convert(source, external_decisions, output_dir=tmp)
        except RuntimeError as exc:
            raised = True
            assert_true("未対応領域" in str(exc), "Error should mention the unsupported region")
        assert_true(raised, "External-share mode must block output when unsupported features are present")

        analysis_processor = WordPrivacyProcessor()
        analysis_decisions = analysis_processor.scan(source)
        result = analysis_processor.convert(source, analysis_decisions, output_dir=tmp)
        assert_true(
            any("未対応領域内の匿名化は保証しない" in warning for warning in result.warnings),
            "Analysis mode must warn with the exact required phrase instead of blocking",
        )
        assert_true("090-1111-2222" not in Document(result.output_path).paragraphs[0].text, "High-confidence candidate should still convert")


def test_candidate_requires_review_threshold() -> None:
    high_confidence = replace(
        WordCandidate(
            candidate_id="a",
            category="氏名",
            story_type="body",
            section_index=None,
            container_type="paragraph",
            paragraph_index=0,
            table_index=None,
            row_index=None,
            cell_index=None,
            location_id="loc",
            char_start=0,
            char_end=2,
            text="x",
            detection_rule="test",
            confidence=0.75,
            source="paragraph",
        )
    )
    low_confidence = replace(high_confidence, candidate_id="b", confidence=0.74)
    weak_hyperlink = replace(high_confidence, candidate_id="c", confidence=0.1, source="hyperlink_target")

    assert_true(not candidate_requires_review(high_confidence), "Confidence at the threshold should not require review")
    assert_true(candidate_requires_review(low_confidence), "Confidence just below the threshold should require review")
    assert_true(not candidate_requires_review(weak_hyperlink), "hyperlink_target must never be flagged by this guard")


def test_find_residual_text_detects_company_in_app_properties() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)
        inject_company_into_app_properties(source, "株式会社未来会議")

        candidate = WordCandidate(
            candidate_id="synthetic_app_props",
            category="会社名",
            story_type="body",
            section_index=None,
            container_type="paragraph",
            paragraph_index=0,
            table_index=None,
            row_index=None,
            cell_index=None,
            location_id="loc",
            char_start=0,
            char_end=8,
            text="株式会社未来会議",
            detection_rule="test",
            confidence=0.9,
            source="paragraph",
        )
        decision = WordReplacementDecision(candidate=candidate, enabled=True, replacement="法人001")

        residual = _find_residual_text(source, [decision])
        assert_true("docProps/app.xml" in residual, "Company text injected into app.xml should be detected")
        assert_true(residual["docProps/app.xml"] == {"会社名"}, "Detected category should match the candidate's category")


def test_residual_text_in_internal_xml_blocks_external_output() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)
        inject_company_into_app_properties(source, "株式会社未来会議")

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source, options=ProcessingOptions(mode="external"))
        _enable_review_required(decisions)
        # python-docx's own default template already contains customXml parts and a
        # styles.xml false-positive (see below), which independently trips the
        # pre-existing (step 6) unsupported-features guard for *any* document. That's
        # a separate, out-of-scope issue from this step's residual-XML guard, so it's
        # neutralized here to isolate the guard actually under test.
        processor.inventory = replace(processor.inventory, unsupported_features=())

        before = set(tmp.iterdir())
        raised = False
        try:
            processor.convert(source, decisions, output_dir=tmp)
        except RuntimeError as exc:
            raised = True
            assert_true("残存" in str(exc) and "docProps/app.xml" in str(exc), "Error should identify the residual XML part")
        assert_true(raised, "Residual text in internal XML must block external-share output")
        assert_true(set(tmp.iterdir()) == before, "No output file should be left behind when blocked")


def test_residual_text_in_internal_xml_warns_in_analysis_mode_without_leaking_original_text() -> None:
    with tempfile.TemporaryDirectory(prefix="word_replacement_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_replacement_fixture(source)
        inject_company_into_app_properties(source, "株式会社未来会議")

        processor = WordPrivacyProcessor()
        decisions = processor.scan(source)
        _enable_review_required(decisions)
        result = processor.convert(source, decisions, output_dir=tmp)

        assert_true(result.output_path.exists(), "Analysis mode must still produce output despite the residual warning")
        residual_warnings = [w for w in result.warnings if "残存" in w]
        assert_true(len(residual_warnings) == 1, "Exactly one residual warning should be recorded")
        assert_true("docProps/app.xml" in residual_warnings[0] and "会社名" in residual_warnings[0], "Warning should name the part and category")
        assert_true(
            all("株式会社未来会議" not in warning for warning in result.warnings),
            "Warnings must never contain the raw leaked original text",
        )
