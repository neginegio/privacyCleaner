from __future__ import annotations

import csv
import json
import re
import sys
import tempfile
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from excel_privacy_cleaner.word_processor import WordCandidate, candidates_for_inventory, extract_word_structure  # noqa: E402


ROOT = Path(__file__).resolve().parents[1]
DATASET_CONFIG = ROOT / "config" / "evaluation" / "word_dataset_split_v1.json"
BANNED_GENERATION_TOKENS = (
    "ground_truth",
    "truth_id",
    "正解",
    "未来会議",
    "山田",
    "佐藤",
    "星雲商事",
    "北斗物流",
    "白樺",
    "鈴木花子",
    "青葉ビル",
)
BANNED_FIXED_LOCATION_PATTERNS = (
    r"paragraph_index\s*==\s*\d+",
    r"table_index\s*==\s*\d+",
    r"row_index\s*==\s*\d+",
    r"cell_index\s*==\s*\d+",
)


@dataclass(frozen=True)
class ExpectedTruth:
    category: str
    marker: str
    occurrence: int = 0
    source: str = "paragraph"


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_development_docx(path: Path) -> None:
    document = Document()
    document.core_properties.author = "山田 太郎"
    document.core_properties.last_modified_by = "株式会社未来会議"
    document.core_properties.title = "開発用 Word Phase 2"
    document.core_properties.comments = "連絡先 taro.yamada@example.jp"

    paragraph = document.add_paragraph()
    paragraph.add_run("株式会社")
    paragraph.add_run("未来会議")
    paragraph.add_run(" の担当は")
    paragraph.add_run("山田")
    paragraph.add_run(" 太郎")
    paragraph.add_run("です。")
    document.add_paragraph("送付先住所は東京都千代田区丸の内1-1-1です。")
    document.add_paragraph("電話番号は090-1111-2222、メールはtaro.yamada@example.jpです。")
    document.add_paragraph("振込先はみらい銀行 東京支店です。")
    document.add_paragraph("株式会社未来会議は本文に再登場します。")
    document.add_paragraph("会社という一般語、銀行という一般語、123-4567、not-an-email@example は対象外です。")
    document.add_paragraph("この文章は会社名に似ていますが、会社の方針という一般文です。")
    linked = document.add_paragraph("リンク先も確認: ")
    add_hyperlink(linked, "会社サイト", "https://example.invalid/contact?email=contact@example.jp&user=yamada")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "会社"
    table.cell(0, 1).text = "氏名"
    table.cell(1, 0).text = "株式会社未来会議"
    table.cell(1, 1).text = "佐藤 花子"

    section = document.sections[0]
    section.header.paragraphs[0].text = "ヘッダー 株式会社未来会議"
    section.footer.paragraphs[0].text = "フッター 山田 太郎"
    document.save(path)


EXPECTED_TRUTHS = (
    ExpectedTruth("会社名", "株式会社未来会議", 0),
    ExpectedTruth("氏名", "山田 太郎", 0),
    ExpectedTruth("住所", "東京都千代田区丸の内1-1-1", 0),
    ExpectedTruth("電話番号", "090-1111-2222", 0),
    ExpectedTruth("メールアドレス", "taro.yamada@example.jp", 0),
    ExpectedTruth("銀行名", "みらい銀行", 0),
    ExpectedTruth("会社名", "株式会社未来会議", 1),
    ExpectedTruth("メールアドレス", "contact@example.jp", 0, source="hyperlink_target"),
    ExpectedTruth("氏名", "yamada", 0, source="hyperlink_target"),
    ExpectedTruth("会社名", "株式会社未来会議", 2),
    ExpectedTruth("氏名", "佐藤 花子", 0),
    ExpectedTruth("会社名", "株式会社未来会議", 3),
    ExpectedTruth("氏名", "山田 太郎", 1),
    ExpectedTruth("氏名", "山田 太郎", 0, source="document_property"),
    ExpectedTruth("会社名", "株式会社未来会議", 0, source="document_property"),
    ExpectedTruth("メールアドレス", "taro.yamada@example.jp", 0, source="document_property"),
)


def build_ground_truth(candidates: tuple[WordCandidate, ...]) -> list[dict[str, object]]:
    truths: list[dict[str, object]] = []
    for index, expected in enumerate(EXPECTED_TRUTHS, start=1):
        matches = [
            candidate
            for candidate in candidates
            if candidate.category == expected.category
            and candidate.text == expected.marker
            and candidate.source == expected.source
        ]
        assert_true(len(matches) > expected.occurrence, f"Expected marker not found before evaluation: {expected}")
        candidate = matches[expected.occurrence]
        truths.append(
            {
                "truth_id": f"WORDDEV-{index:03d}",
                "category": expected.category,
                "location_id": candidate.location_id,
                "char_start": candidate.char_start,
                "char_end": candidate.char_end,
                "text": candidate.text,
                "source": candidate.source,
            }
        )
    return truths


def evaluate_candidates(candidates: tuple[WordCandidate, ...], truths: list[dict[str, object]]) -> tuple[dict[str, object], dict[str, dict[str, int]]]:
    truth_keys = {
        (str(truth["location_id"]), str(truth["category"]), int(truth["char_start"]), int(truth["char_end"]))
        for truth in truths
    }
    candidate_keys = {
        (candidate.location_id, candidate.category, candidate.char_start, candidate.char_end)
        for candidate in candidates
    }
    matched_truth_keys = truth_keys & candidate_keys
    true_positive_keys = candidate_keys & truth_keys
    false_positive_keys = candidate_keys - truth_keys
    missed_truth_keys = truth_keys - candidate_keys
    metrics = {
        "ground_truth_count": len(truth_keys),
        "candidate_count": len(candidate_keys),
        "matched_ground_truth_count": len(matched_truth_keys),
        "missed_ground_truth_count": len(missed_truth_keys),
        "true_positive_candidate_count": len(true_positive_keys),
        "false_positive_candidate_count": len(false_positive_keys),
        "recall": len(matched_truth_keys) / len(truth_keys) if truth_keys else 0.0,
        "candidate_precision": len(true_positive_keys) / len(candidate_keys) if candidate_keys else 0.0,
        "exact_span_match_count": len(matched_truth_keys),
    }
    by_category: dict[str, dict[str, int]] = defaultdict(lambda: {"正解": 0, "検出": 0, "見逃し": 0, "TP候補": 0, "FP候補": 0})
    for _location, category, _start, _end in truth_keys:
        by_category[category]["正解"] += 1
    for _location, category, _start, _end in candidate_keys:
        by_category[category]["検出"] += 1
    for _location, category, _start, _end in missed_truth_keys:
        by_category[category]["見逃し"] += 1
    for _location, category, _start, _end in true_positive_keys:
        by_category[category]["TP候補"] += 1
    for _location, category, _start, _end in false_positive_keys:
        by_category[category]["FP候補"] += 1
    return metrics, dict(sorted(by_category.items()))


def write_dataset_config(path: Path, docx_path: Path, source_sha256: str) -> None:
    payload = {
        "dataset_version": "word_dataset_split_v1",
        "development": {"synthetic_phase2_docx": {"path": str(docx_path), "sha256": source_sha256}},
        "validation": {"reserved": True, "note": "Do not use validation data for Phase 2 rule tuning."},
        "holdout": {"reserved": True, "note": "Do not inspect or use holdout data during implementation."},
        "leakage_guard": "Candidate generation code must not read this file.",
    }
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(json.dumps(payload, ensure_ascii=False, indent=2) + "\n", encoding="utf-8")


def write_ground_truth_csv(path: Path, truths: list[dict[str, object]]) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    with path.open("w", encoding="utf-8-sig", newline="") as output:
        writer = csv.DictWriter(output, fieldnames=["truth_id", "category", "location_id", "char_start", "char_end", "text", "source"])
        writer.writeheader()
        for truth in truths:
            writer.writerow(truth)


def contamination_guard() -> list[str]:
    source_path = ROOT / "src" / "excel_privacy_cleaner" / "word_processor.py"
    text = source_path.read_text(encoding="utf-8")
    failures: list[str] = []
    for token in BANNED_GENERATION_TOKENS:
        if token in text:
            failures.append(f"candidate_generation_contains:{token}")
    for pattern in BANNED_FIXED_LOCATION_PATTERNS:
        if re.search(pattern, text):
            failures.append(f"fixed_location_pattern:{pattern}")
    if "word_dataset_split_v1" in text or "holdout" in text:
        failures.append("candidate_generation_references_dataset_split_or_holdout")
    return failures


def run_phase2_1_unit_tests() -> None:
    with tempfile.TemporaryDirectory(prefix="word_phase2_1_units_") as tmpdir:
        docx_path = Path(tmpdir) / "word_phase2_1_units.docx"
        document = Document()
        document.add_paragraph("配送担当は北斗物流株式会社です。")
        document.add_paragraph("共同研究先は合同会社青葉研究所です。")
        document.add_paragraph("法人格のみ: 株式会社")
        document.add_paragraph("法人格を含まない組織として白樺ラボを記載します。")
        document.add_paragraph("会社の方針を確認します。")
        document.add_paragraph("連絡者は鈴木花子です。")
        document.add_paragraph("一般文章として南側の入口、森を抜ける話です。")
        document.add_paragraph("住所: 京都府京都市下京区烏丸通1-2 青葉ビル5階")
        document.add_paragraph("所在地: 東京都千代田区丸の内1-1-1 5階")
        document.add_paragraph("市区町村の制度について説明します。")
        document.add_paragraph("東京都の政策を確認します。")
        document.add_paragraph("送付先住所は東京都千代田区丸の内1-1-1です。")
        document.save(docx_path)

        inventory = extract_word_structure(docx_path)
        candidates = candidates_for_inventory(inventory)
        by_text_category = {(candidate.text, candidate.category) for candidate in candidates}
        candidate_texts = {candidate.text for candidate in candidates}

        assert_true(("北斗物流株式会社", "会社名") in by_text_category, "Suffix company should keep the exact company span")
        assert_true(("合同会社青葉研究所", "会社名") in by_text_category, "Prefix company should still be detected")
        assert_true(("配送担当", "会社名") not in by_text_category, "Particle-leading context should not be included as company")
        assert_true(("共同研究先", "会社名") not in by_text_category, "Context label should not become a company")
        assert_true(("株式会社", "会社名") not in by_text_category, "Designator alone should not be a company")
        assert_true(("北斗物流株式会社", "氏名") not in by_text_category, "Corporate names should not remain as person candidates")
        assert_true(("白樺ラボ", "会社名") in by_text_category, "Contextual organization should be detected with clear evidence")
        assert_true(("会社", "会社名") not in by_text_category, "Generic company word should not be a candidate")
        assert_true(("鈴木花子", "氏名") in by_text_category, "Person role label should detect a no-space Japanese name")
        assert_true("南側" not in candidate_texts and "森" not in candidate_texts, "Name-like general words should not be candidates")
        assert_true(("京都府京都市下京区烏丸通1-2 青葉ビル5階", "住所") in by_text_category, "Address should include building name")
        assert_true(("東京都千代田区丸の内1-1-1 5階", "住所") in by_text_category, "Address should include floor")
        assert_true(("市区町村の制度", "住所") not in by_text_category, "General municipality sentence should not be an address")
        assert_true(all(not (candidate.text == "東京都の政策" and candidate.category == "住所") for candidate in candidates), "General Tokyo policy sentence should not be an address")
        assert_true(("東京都千代田区丸の内1-1-1", "住所") in by_text_category, "Normal address should still be detected")


def main() -> int:
    run_phase2_1_unit_tests()
    with tempfile.TemporaryDirectory(prefix="word_phase2_dev_") as tmpdir:
        tmp = Path(tmpdir)
        docx_path = tmp / "word_phase2_development.docx"
        ground_truth_path = tmp / "word_phase2_development_ground_truth.csv"
        create_development_docx(docx_path)
        inventory = extract_word_structure(docx_path)
        candidates = candidates_for_inventory(inventory)
        truths = build_ground_truth(candidates)
        metrics, by_category = evaluate_candidates(candidates, truths)
        write_dataset_config(tmp / "word_dataset_split_v1.runtime.json", docx_path, inventory.source_sha256)
        write_ground_truth_csv(ground_truth_path, truths)

        assert_true(metrics["ground_truth_count"] == 16, "Development ground truth count should be fixed")
        assert_true(metrics["matched_ground_truth_count"] == metrics["ground_truth_count"], "All development truths should match exactly")
        assert_true(metrics["missed_ground_truth_count"] == 0, "No development truths should be missed")
        assert_true(metrics["false_positive_candidate_count"] == 0, f"Unexpected false positives: {metrics}")
        assert_true(metrics["recall"] == 1.0, "Recall should be 1.0 for fixed development fixture")
        assert_true(metrics["candidate_precision"] == 1.0, "Candidate precision should be 1.0 for fixed development fixture")

        categories = {candidate.category for candidate in candidates}
        for category in {"会社名", "氏名", "住所", "電話番号", "メールアドレス", "銀行名"}:
            assert_true(category in categories, f"Missing category: {category}")
        assert_true(any(candidate.container_type == "table_cell" for candidate in candidates), "Table candidates should exist")
        assert_true(any(candidate.story_type == "header" for candidate in candidates), "Header candidates should exist")
        assert_true(any(candidate.story_type == "footer" for candidate in candidates), "Footer candidates should exist")
        assert_true(any(candidate.text == "株式会社未来会議" and candidate.affected_run_indices == (0, 1) for candidate in candidates), "Split company should map to multiple runs")
        assert_true(any(candidate.text == "山田 太郎" and candidate.affected_run_indices == (3, 4) for candidate in candidates), "Split name should map to multiple runs")
        assert_true(len({candidate.location_id for candidate in candidates if candidate.text == "株式会社未来会議"}) >= 4, "Duplicate company text should have distinct locations")
        assert_true(all(candidate.text not in {"会社", "銀行"} for candidate in candidates), "Generic words should not be candidates")
        assert_true(all(candidate.text != "123-4567" for candidate in candidates), "Phone-like short number should not be a candidate")
        assert_true(all(candidate.text != "not-an-email@example" for candidate in candidates), "Incomplete email should not be a candidate")
        assert_true(all(candidate.char_end > candidate.char_start for candidate in candidates), "Candidate spans should be valid")
        assert_true(not contamination_guard(), f"Contamination guard failed: {contamination_guard()}")

    print("word_candidate_detection_tests=passed")
    print("metrics", json.dumps(metrics, ensure_ascii=False, sort_keys=True))
    print("by_category", json.dumps(by_category, ensure_ascii=False, sort_keys=True))
    return 0


def test_word_phase2_1_unit_rules() -> None:
    run_phase2_1_unit_tests()


def test_word_candidate_detection_pipeline() -> None:
    assert main() == 0


if __name__ == "__main__":
    raise SystemExit(main())
