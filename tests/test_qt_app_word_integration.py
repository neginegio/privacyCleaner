from __future__ import annotations

import os
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

os.environ.setdefault("QT_QPA_PLATFORM", "offscreen")

from docx import Document  # noqa: E402
from PySide6.QtCore import Qt  # noqa: E402
from PySide6.QtWidgets import QApplication, QMessageBox  # noqa: E402

from excel_privacy_cleaner.qt_app import ExcelPrivacyCleanerWindow  # noqa: E402
from excel_privacy_cleaner.word_processor import WordPrivacyProcessor  # noqa: E402


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def create_gui_fixture(path: Path) -> None:
    document = Document()
    document.add_paragraph("連絡先電話は090-1111-2222です。")
    document.add_paragraph("佐藤 花子")
    document.save(path)


def create_overlapping_candidates_fixture(path: Path) -> None:
    # Company detector matches "株式会社未来会議" [0:8); the bare-name heuristic
    # matches "未来会議　細谷" [4:11) -- these overlap at [4:8), reproducing a
    # real conflict seen with actual documents.
    document = Document()
    document.add_paragraph("株式会社未来会議　細谷")
    document.save(path)


def _silence_message_boxes(monkeypatch) -> list[tuple[str, str]]:
    calls: list[tuple[str, str]] = []
    for method in ("information", "warning", "critical", "question"):
        def make_stub(name: str):
            def stub(*_args, **_kwargs):
                calls.append((name, ""))
                return QMessageBox.Yes if name == "question" else None
            return stub
        monkeypatch.setattr(QMessageBox, method, staticmethod(make_stub(method)))
    return calls


def test_word_scan_and_convert_via_gui(monkeypatch) -> None:
    app = QApplication.instance() or QApplication([])
    calls = _silence_message_boxes(monkeypatch)

    with tempfile.TemporaryDirectory(prefix="qt_app_word_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_gui_fixture(source)

        window = ExcelPrivacyCleanerWindow()
        try:
            window.set_source(source)
            assert_true(isinstance(window.processor, WordPrivacyProcessor), "Selecting a .docx should create a WordPrivacyProcessor")
            assert_true(window.is_word_source(), "is_word_source() should be True after selecting a .docx")

            window.scan_file()
            assert_true(len(window.word_decisions) == 2, "Fixture should yield exactly two candidates")
            assert_true(window.table.rowCount() == 2, "Table should show one row per candidate")
            assert_true(len(window.findings) == 2, "Adapter should produce one Finding per WordReplacementDecision")

            phone_row = next(row for row, finding in enumerate(window.findings) if finding.entity_type == "電話番号")
            name_row = next(row for row, finding in enumerate(window.findings) if finding.entity_type == "氏名")
            assert_true(window.word_decisions[phone_row].enabled, "High-confidence phone candidate should default to enabled")
            assert_true(not window.word_decisions[name_row].enabled, "Low-confidence bare name candidate should default to disabled (review-required)")

            # Column 4 (entity_type) must not be user-editable for Word rows,
            # since WordCandidate.category is immutable.
            entity_item = window.table.item(name_row, 4)
            assert_true(entity_item is not None and not (entity_item.flags() & Qt.ItemIsEditable), "Word rows must not allow editing the category column")

            # Simulate a reviewer approving the review-required candidate via the checkbox.
            checkbox_item = window.table.item(name_row, 0)
            assert_true(checkbox_item is not None, "Checkbox item should exist")
            checkbox_item.setCheckState(Qt.Checked)

            window.convert_file()

            assert_true(any(name == "information" for name, _ in calls), "A success message box should have been shown")
            assert_true(window.history.count() == 1, "Conversion should append one history entry")
        finally:
            window.processor.cleanup()
            window.close()
            app.processEvents()


def test_word_exclude_checkbox_resolves_overlap_and_review_required_deadlock(monkeypatch) -> None:
    app = QApplication.instance() or QApplication([])
    calls = _silence_message_boxes(monkeypatch)

    with tempfile.TemporaryDirectory(prefix="qt_app_word_overlap_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "fixture.docx"
        create_overlapping_candidates_fixture(source)

        window = ExcelPrivacyCleanerWindow()
        try:
            window.set_source(source)
            window.scan_file()
            assert_true(len(window.word_decisions) == 2, "Fixture should yield the company and the overlapping name candidate")

            name_row = next(row for row, finding in enumerate(window.findings) if finding.entity_type == "氏名")
            assert_true(not window.word_decisions[name_row].enabled, "Overlapping low-confidence name should default to disabled")

            # Enabling both (company auto-enabled by default, name checked here
            # too) must be blocked by the overlap guard.
            window.table.item(name_row, 0).setCheckState(Qt.Checked)
            window.convert_file()
            assert_true(any(name == "critical" for name, _ in calls), "Enabling both overlapping candidates must raise an error, not silently convert")
            calls.clear()

            # Uncheck it again (still just "unresolved", not "reviewed") --
            # must still be blocked, now by the review-required guard.
            window.table.item(name_row, 0).setCheckState(Qt.Unchecked)
            window.convert_file()
            assert_true(any(name == "critical" for name, _ in calls), "Leaving a review-required candidate merely unchecked must still block")
            calls.clear()

            # Check the "変換しない" (don't convert) checkbox for the name row --
            # this is the actual fix for the deadlock.
            window.table.item(name_row, 1).setCheckState(Qt.Checked)
            assert_true(window.word_decisions[name_row].excluded, "Decision should now be marked excluded")
            assert_true(not window.word_decisions[name_row].enabled, "Excluded decision must stay disabled")
            assert_true(
                window.table.item(name_row, 0).checkState() == Qt.Unchecked,
                "Checking 変換しない must clear 変換する for the same row",
            )

            result_holder: list[object] = []
            original_convert = window.processor.convert

            def capture_convert(*args, **kwargs):
                result = original_convert(*args, **kwargs)
                result_holder.append(result)
                return result

            monkeypatch.setattr(window.processor, "convert", capture_convert)
            window.convert_file()
            assert_true(bool(result_holder), "convert() should have succeeded this time")
            assert_true(any(name == "information" for name, _ in calls), "A success message box should have been shown after resolving via exclude")

            output_text = Document(result_holder[0].output_path).paragraphs[0].text
            # The excluded name candidate's span [4:11) overlaps the enabled
            # company candidate's span [0:8): "未来会議" falls inside both, so
            # it's consumed by the company replacement regardless of the name
            # decision being excluded. Only "細谷" (outside the company span)
            # is actually left untouched by excluding the name candidate.
            assert_true("細谷" in output_text, "The part of the excluded candidate outside the company span must remain untouched")
            assert_true("株式会社未来会議" not in output_text, "The company candidate should still have been converted")
        finally:
            window.processor.cleanup()
            window.close()
            app.processEvents()
