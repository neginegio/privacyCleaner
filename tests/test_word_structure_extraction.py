from __future__ import annotations

import shutil
import sys
import tempfile
import zipfile
from pathlib import Path
from xml.sax.saxutils import escape

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE as RT

sys.path.insert(0, str(Path(__file__).resolve().parents[1] / "src"))

from excel_privacy_cleaner.word_processor import UnsupportedWordDocumentError, extract_word_structure  # noqa: E402


def assert_true(condition: bool, message: str) -> None:
    if not condition:
        raise AssertionError(message)


def sha256(path: Path) -> str:
    import hashlib

    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def add_hyperlink(paragraph, text: str, url: str) -> None:
    rel_id = paragraph.part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    run = OxmlElement("w:r")
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def create_fixture(path: Path) -> None:
    document = Document()
    props = document.core_properties
    props.author = "Test Author"
    props.last_modified_by = "Test Editor"
    props.title = "Word Phase 1 Fixture"
    props.subject = "Structure extraction"
    props.keywords = "privacy,word"
    props.category = "test"
    props.comments = "Synthetic document only"

    paragraph = document.add_paragraph()
    paragraph.add_run("株式会社")
    paragraph.add_run("未来会議")
    paragraph.add_run(" の 山田 太郎")

    document.add_paragraph("住所は東京都千代田区丸の内1-1-1です。")
    document.add_paragraph("電話 090-1111-2222 / mail taro@example.jp")
    document.add_paragraph("同じ会社名 株式会社未来会議 が本文にもあります。")
    hyperlink_paragraph = document.add_paragraph("リンク: ")
    add_hyperlink(hyperlink_paragraph, "株式会社未来会議サイト", "https://example.com/company/mirai")

    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "会社名"
    table.cell(0, 1).text = "氏名"
    table.cell(1, 0).text = "株式会社未来会議"
    name_cell = table.cell(1, 1)
    name_cell.text = ""
    name_cell.paragraphs[0].add_run("山田")
    name_cell.paragraphs[0].add_run(" 太郎")
    name_cell.add_paragraph("銀行名: みらい銀行")

    section = document.sections[0]
    section.header.paragraphs[0].text = "ヘッダー 株式会社未来会議"
    section.footer.paragraphs[0].text = "フッター 山田 太郎"
    header_table = section.header.add_table(rows=1, cols=1, width=1)
    header_table.cell(0, 0).text = "ヘッダー表 株式会社未来会議"

    document.save(path)
    inject_unsupported_parts(path)


def inject_unsupported_parts(path: Path) -> None:
    injected = (
        "<w:p>"
        "<w:r><w:fldChar w:fldCharType=\"begin\"/></w:r>"
        "<w:r><w:instrText> MERGEFIELD CustomerName </w:instrText></w:r>"
        "<w:r><w:fldChar w:fldCharType=\"end\"/></w:r>"
        "</w:p>"
        "<w:p><w:ins><w:r><w:t>変更履歴テキスト</w:t></w:r></w:ins></w:p>"
        "<w:p><w:r><w:pict><v:shape><v:textbox><w:txbxContent>"
        "<w:p><w:r><w:t>テキストボックス内文字</w:t></w:r></w:p>"
        "</w:txbxContent></v:textbox></v:shape></w:pict></w:r></w:p>"
    )
    replacements = {
        "word/comments.xml": (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:comments xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
            "<w:comment w:id=\"0\"><w:p><w:r><w:t>コメント内文字</w:t></w:r></w:p></w:comment>"
            "</w:comments>"
        ).encode("utf-8"),
        "word/footnotes.xml": (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:footnotes xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
            "<w:footnote w:id=\"1\"><w:p><w:r><w:t>脚注文字</w:t></w:r></w:p></w:footnote>"
            "</w:footnotes>"
        ).encode("utf-8"),
        "word/endnotes.xml": (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:endnotes xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
            "<w:endnote w:id=\"1\"><w:p><w:r><w:t>文末脚注文字</w:t></w:r></w:p></w:endnote>"
            "</w:endnotes>"
        ).encode("utf-8"),
        "customXml/privacyCleanerPhase1Unsupported.xml": f"<root>{escape('カスタムXML文字')}</root>".encode("utf-8"),
    }
    tmp_path = path.with_name(path.name + ".tmp")
    with zipfile.ZipFile(path) as source_archive, zipfile.ZipFile(tmp_path, "w", compression=zipfile.ZIP_DEFLATED) as target_archive:
        for info in source_archive.infolist():
            if info.filename in replacements:
                continue
            data = source_archive.read(info.filename)
            if info.filename == "word/document.xml":
                document_xml = data.decode("utf-8")
                data = document_xml.replace("</w:body>", injected + "</w:body>").encode("utf-8")
            target_archive.writestr(info, data)
        for name, data in replacements.items():
            target_archive.writestr(name, data)
    tmp_path.replace(path)


def main() -> int:
    with tempfile.TemporaryDirectory(prefix="word_structure_test_") as tmpdir:
        tmp = Path(tmpdir)
        source = tmp / "synthetic_word_phase1.docx"
        create_fixture(source)
        before_sha = sha256(source)
        inventory = extract_word_structure(source)
        after_sha = sha256(source)

        assert_true(before_sha == after_sha == inventory.source_sha256, "Reading structure should not modify the DOCX")

        body_paragraphs = [item for item in inventory.paragraphs if item.story_type == "body" and item.container_type == "paragraph"]
        assert_true(any("株式会社未来会議 の 山田 太郎" == item.text for item in body_paragraphs), "Body paragraph should be extracted")

        target_runs = [item for item in inventory.runs if item.paragraph_text == "株式会社未来会議 の 山田 太郎"]
        assert_true([(run.text, run.char_start, run.char_end) for run in target_runs[:3]] == [("株式会社", 0, 4), ("未来会議", 4, 8), (" の 山田 太郎", 8, 16)], "Run offsets should match paragraph text")

        table_runs = [item for item in inventory.runs if item.container_type == "table_cell"]
        assert_true(any(run.table_index == 0 and run.row_index == 1 and run.cell_index == 0 and run.text == "株式会社未来会議" for run in table_runs), "Table cell text should be extracted")
        assert_true(any(run.cell_paragraph_index == 1 and run.text == "銀行名: みらい銀行" for run in table_runs), "Second paragraph in table cell should be extracted")

        assert_true(any(run.story_type == "header" and "ヘッダー" in run.text for run in inventory.runs), "Header text should be extracted")
        assert_true(any(run.story_type == "footer" and "フッター" in run.text for run in inventory.runs), "Footer text should be extracted")
        assert_true(any(run.story_type == "header" and run.container_type == "table_cell" and "ヘッダー表" in run.text for run in inventory.runs), "Header table text should be extracted")

        duplicate_locations = [run.location_id for run in inventory.runs if "株式会社未来会議" in run.text]
        assert_true(len(duplicate_locations) >= 4, "Duplicate text should appear in multiple locations")
        assert_true(len(duplicate_locations) == len(set(duplicate_locations)), "Duplicate text locations should be unique")

        assert_true(any(link.text == "株式会社未来会議サイト" and link.url == "https://example.com/company/mirai" for link in inventory.hyperlinks), "Hyperlink display text and URL should be extracted")
        assert_true(inventory.document_properties.author == "Test Author", "Author property should be extracted")
        assert_true(inventory.document_properties.last_modified_by == "Test Editor", "Last modified by property should be extracted")
        assert_true(inventory.document_properties.title == "Word Phase 1 Fixture", "Title property should be extracted")

        feature_types = {feature.feature_type for feature in inventory.unsupported_features}
        for expected in {
            "textbox",
            "comments",
            "footnotes",
            "endnotes",
            "tracked_changes",
            "complex_field_code",
            "custom_xml",
            "external_link",
            "text_holding_xml_part",
        }:
            assert_true(expected in feature_types, f"Unsupported inventory should include {expected}")

        docm = tmp / "synthetic_word_phase1.docm"
        shutil.copyfile(source, docm)
        rejected = False
        try:
            extract_word_structure(docm)
        except UnsupportedWordDocumentError:
            rejected = True
        assert_true(rejected, ".docm should be rejected")

    print("word_structure_extraction_tests=passed")
    return 0


def test_word_structure_extraction() -> None:
    assert main() == 0


if __name__ == "__main__":
    raise SystemExit(main())
